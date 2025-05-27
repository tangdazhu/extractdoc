import os
import logging
from .libreoffice_converter import convert_to_pdf as lo_convert_to_pdf

logger = logging.getLogger(__name__)

def convert_word_to_pdf(input_word_path, output_pdf_path):
    """
    Converts a Word document (.doc, .docx) to a PDF file using LibreOffice.

    Args:
        input_word_path (str): Absolute path to the input Word document.
        output_pdf_path (str): Absolute path where the PDF should be saved.
                               Note: libreoffice_converter will handle placing the
                               file in the correct directory based on its --outdir
                               parameter and naming it based on the input.
                               This function primarily ensures the call and then
                               verifies/moves if the final output_pdf_path name is specific.

    Returns:
        tuple: (bool_success, pdf_output_path_or_error_msg, original_pdf_filename_or_None)
               - bool_success: True if conversion was successful, False otherwise.
               - pdf_output_path_or_error_msg: Absolute path to the converted PDF if successful,
                                               or an error message string if failed.
               - original_pdf_filename_or_None: The original filename of the PDF as created by LibreOffice.
    """
    logger.info(f"Attempting to convert Word '{input_word_path}' to PDF using LibreOffice converter.")

    if not os.path.exists(input_word_path):
        return False, f"Input Word file not found: {input_word_path}", None

    # The output_dir for LibreOffice should be the directory part of output_pdf_path
    output_dir = os.path.dirname(output_pdf_path)
    
    # Ensure output directory exists
    if not os.path.isdir(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
            logger.info(f"Created output directory: {output_dir}")
        except OSError as e:
            return False, f"Failed to create output directory {output_dir}: {e}", None

    success, lo_result_path_or_msg, lo_original_filename = lo_convert_to_pdf(input_word_path, output_dir)

    if success:
        # LibreOffice creates a file like input_word_filename.pdf in output_dir.
        # If the desired output_pdf_path is different (e.g., has a unique ID in its name),
        # we might need to rename/move.
        # lo_result_path_or_msg is the actual path of the PDF created by LibreOffice.
        
        # If the actual path from LibreOffice is not the desired final output_pdf_path, move it.
        if lo_result_path_or_msg != output_pdf_path:
            try:
                # Ensure the target path (output_pdf_path) doesn't already exist from a previous attempt or other file
                if os.path.exists(output_pdf_path):
                    logger.warning(f"Target output PDF path {output_pdf_path} already exists. Overwriting.")
                    os.remove(output_pdf_path)
                
                os.rename(lo_result_path_or_msg, output_pdf_path)
                logger.info(f"Successfully moved/renamed '{lo_result_path_or_msg}' to '{output_pdf_path}'")
                return True, output_pdf_path, os.path.basename(output_pdf_path) # Return new name
            except Exception as e_move:
                logger.error(f"LibreOffice Word to PDF conversion succeeded, but failed to move/rename PDF from '{lo_result_path_or_msg}' to '{output_pdf_path}': {e_move}", exc_info=True)
                # Return success as true, but with the path where LO left the file, and an error message indicating move failure.
                # Or, consider this a failure if the final path is critical. For now, let's say it's a partial success.
                return True, lo_result_path_or_msg, f"Conversion OK, but move to final path failed: {e_move}. File at: {lo_result_path_or_msg}"
        else:
            # The file is already where we want it.
            return True, output_pdf_path, lo_original_filename
    else:
        return False, lo_result_path_or_msg, None

if __name__ == '__main__':
    # Example Usage (for testing this script directly)
    # You would need a test Word document (e.g., test_document.docx)
    # and libreoffice_converter.py in the same directory or adjust import path.
    
    logging.basicConfig(level=logging.DEBUG)
    logger.info("Testing Word to PDF Converter (via LibreOffice)")

    # Create a dummy docx for testing if python-docx is available
    # Note: libreoffice_converter.py itself has a more robust dummy file creation.
    # This is a simplified one.
    test_input_dir = "test_io_word_to_pdf"
    test_output_dir = os.path.join(test_input_dir, "output")
    os.makedirs(test_output_dir, exist_ok=True)
    
    dummy_docx_path = os.path.join(test_input_dir, "my_test_word_doc.docx")
    desired_pdf_path = os.path.join(test_output_dir, "my_converted_word_doc_final.pdf")

    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading('Test Document for Word to PDF', 0)
        doc.add_paragraph('This is a test paragraph.')
        doc.save(dummy_docx_path)
        print(f"Created dummy DOCX: {dummy_docx_path}")
    except ImportError:
        print("python-docx not installed. Cannot create a dummy .docx for testing. Please create one manually.")
        # Fallback: create a simple text file that soffice might interpret as a doc or error on.
        with open(dummy_docx_path, "w") as f:
            f.write("This is a simple text file, named .docx for testing word to pdf conversion.")
        print(f"Created dummy text file (named .docx): {dummy_docx_path}")
    except Exception as e_create:
        print(f"Error creating dummy docx: {e_create}")
        dummy_docx_path = None


    if dummy_docx_path and os.path.exists(dummy_docx_path):
        abs_input_path = os.path.abspath(dummy_docx_path)
        abs_output_path = os.path.abspath(desired_pdf_path)
        
        print(f"Converting '{abs_input_path}' to '{abs_output_path}'...")
        success, result_msg_or_path, original_name = convert_word_to_pdf(abs_input_path, abs_output_path)

        if success:
            print(f"Word to PDF Conversion successful!")
            print(f"Final PDF at: {result_msg_or_path}")
            if original_name:
                print(f"Original PDF filename by LO (or final if moved): {original_name}")
            if os.path.exists(abs_output_path):
                print(f"VERIFIED: Final PDF exists at desired path: {abs_output_path}")
            else:
                print(f"ERROR: Final PDF DOES NOT exist at desired path: {abs_output_path}, though result was: {result_msg_or_path}")
        else:
            print(f"Word to PDF Conversion failed.")
            print(f"Error/Path: {result_msg_or_path}")
    else:
        print(f"Skipping test, input file {dummy_docx_path} not created or found.")

    # Consider cleanup:
    # import shutil
    # if os.path.exists(test_input_dir):
    #     shutil.rmtree(test_input_dir)
    #     print(f"Cleaned up {test_input_dir}") 