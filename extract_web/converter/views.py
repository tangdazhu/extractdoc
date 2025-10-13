from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login
from .forms import RegistrationForm, AdminUserEditForm, AdminSetPasswordForm  # 更新导入
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.conf import settings
import os
import subprocess  # For running the script
import sys  # <--- ADDED IMPORT FOR SYS
import json  # <--- ADDED IMPORT FOR JSON
import re  # <--- ADDED IMPORT FOR RE
from django.contrib import messages  # 新增导入
from django.http import (
    JsonResponse,
    FileResponse,
    Http404,
    StreamingHttpResponse,
    HttpResponseBadRequest,
)
from django.views.decorators.http import require_POST  # To restrict to POST requests
from django.views.decorators.csrf import csrf_exempt  # <<< Import csrf_exempt
import random
import string
import traceback  # 新增导入 for detailed exception logging
import logging  # 新增导入
import time  # ADDED IMPORT
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement  # For adding content from sub-documents
from docx.oxml.ns import qn
from pathlib import Path  # 新增
from datetime import datetime  # 新增 datetime
from django.urls import reverse
import shutil  # Import shutil earlier as it's used in multiple places
from converter import ppt_pdf_converter  # 新的导入方式
from .pic_file_converter import process_images_to_files  # 导入图片转文件模块
from .excel_pdf_converter import convert_excel_to_pdf  # 导入Excel转换模块
from .txt_to_pdf_converter import convert_txt_to_pdf
from .pdf_to_excel_converter import convert_pdf_to_excel
from .pdf_to_word_converter import convert_pdf_to_word, convert_and_merge_pdfs_to_docx

# Add new imports for PDF to X converters
from .pdf_to_ppt_converter import convert_pdf_to_ppt, convert_and_merge_pdfs_to_pptx
from .pdf_to_txt_converter import convert_pdf_to_txt, convert_and_merge_pdfs_to_txt
from .libreoffice_converter import convert_to_pdf as convert_to_pdf_libreoffice
from .libreoffice_converter import (
    convert_to_pptx as convert_docx_to_pptx_libreoffice,
)  # Added import for DOCX to PPTX
from .word_to_pdf_converter import (
    convert_word_to_pdf,
)  # ADDED: Import for the new Word to PDF converter
from .image_to_pptx import copy_images_to_pptx  # 导入直接图片转PPTX函数
from django.core.exceptions import PermissionDenied  # For security checks
from .speech_processor import (
    transcribe_audio_dashscope,
)  # ADDED: Import for speech transcription
from .text_to_voice import (
    get_predefined_tts_voices,
)  # ADDED: Import for TTS voice list
from .realtime_speech_view import (
    start_realtime_recognition,
    send_audio_data,
    get_recognition_results,
    stop_realtime_recognition,
)  # ADDED: Import for real-time speech recognition

# ADDED: For text extraction from PDF in TTS, make it an optional import
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    logging.warning(
        "pdfplumber is not installed. PDF text extraction for TTS will not work."
    )

# Import the new file handling utility
from .utils.file_handling import (
    ensure_user_directories,
    save_uploaded_file,
    delete_user_data_folder,
    cleanup_temp_files,  # Added import
)
from .utils.request_parsing import parse_conversion_request_params  # Added import
from .services.response_formatters import (
    format_json_response,
    format_error_response,
)  # Added import

# PDF to X Merge Converters
from .pdf_to_word_converter import convert_and_merge_pdfs_to_docx
from .pdf_to_ppt_converter import convert_and_merge_pdfs_to_pptx
from .pdf_to_txt_converter import convert_and_merge_pdfs_to_txt

# Try to import docx2pdf and set a flag
DOCX2PDF_AVAILABLE_IN_VIEW = False
try:
    from docx2pdf import convert as docx_to_pdf_converter_internal

    DOCX2PDF_AVAILABLE_IN_VIEW = True
except ImportError:
    logging.warning(
        "docx2pdf library is not installed. Word to PDF conversion will be skipped for direct docx2pdf method."
    )

    def docx_to_pdf_converter_internal(input_path, output_path):
        logging.error(
            "docx2pdf is not available, cannot convert Word to PDF using this method."
        )
        raise NotImplementedError("docx2pdf is not installed.")


logger = logging.getLogger("converter")  # 获取 logger 实例


# Helper function to generate a unique request ID
def generate_request_id(length=10):
    return "".join(random.choices(string.ascii_lowercase + string.digits, k=length))


# Attempt to import PyPDF2 for PDF merging
try:
    from PyPDF2 import PdfMerger, PdfReader

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning(
        "PyPDF2 library is not installed. Merging multiple PPT/PPTX files into a single PDF will not be available."
    )

# Create your views here.


def index(request):
    """
    Renders the main converter page, passing the curated list of TTS voices
    to the template.
    """
    # This block ensures the static list from text_to_voice.py is always called.
    try:
        tts_voices = get_predefined_tts_voices()
    except Exception as e:
        logging.error(f"Critical error getting predefined TTS voices: {e}")
        tts_voices = []

    context = {"tts_voices": tts_voices}
    return render(request, "converter/index.html", context)


@login_required
def register(request):
    if request.method == "POST":
        form = RegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)

            try:
                # 注册时只创建用户主目录 his_pic/<username>
                # 日期目录将在 process_images_view 中按需创建
                user_main_dir = os.path.join(
                    settings.BASE_DIR, "his_pic", user.username
                )
                os.makedirs(user_main_dir, exist_ok=True)
                logger.info(
                    f"Created main directory for user {user.username} at {user_main_dir}"
                )
            except OSError as e:
                logger.error(
                    f"Error creating main directory for user {user.username}: {e}"
                )

            return redirect("converter:index")
    else:
        form = RegistrationForm()
    return render(request, "registration/register.html", {"form": form})


# Helper to check if user is superuser
def is_superuser(user):
    return user.is_superuser


@login_required
@user_passes_test(is_superuser)
def admin_console_index(request):
    return render(request, "custom_admin/admin_index.html")


@login_required
@user_passes_test(is_superuser)
def admin_user_management(request):
    # Placeholder: Add logic for user CRUD operations here
    users = User.objects.all()
    return render(request, "custom_admin/user_management.html", {"users": users})


@login_required
@user_passes_test(is_superuser)
def admin_file_management(request):
    # Placeholder: Add logic for file management here
    return render(request, "custom_admin/file_management.html")


@login_required
@user_passes_test(is_superuser)
def admin_delete_user(request, user_id):
    if request.method == "POST":
        user_to_delete = get_object_or_404(User, pk=user_id)
        if user_to_delete.is_superuser and not request.user.is_superuser:
            messages.error(request, "您没有权限删除超级管理员用户。")
        elif user_to_delete == request.user:
            messages.error(request, "您不能删除您自己的账户。")
        else:
            username = user_to_delete.username
            # Delete user's data folder using the utility function
            delete_success, delete_message = delete_user_data_folder(username)

            if delete_success:
                # It's possible the folder didn't exist, which is also a form of success for this step.
                # The message from delete_user_data_folder will be informative.
                if (
                    "成功删除" in delete_message or "does not exist" in delete_message
                ):  # Crude check, can be refined
                    messages.success(request, delete_message)
                else:  # Some other non-error message from utility
                    messages.info(request, delete_message)
            else:
                messages.error(request, delete_message)  # Error message from utility

            # Proceed to delete the user object regardless of folder deletion outcome,
            # but maybe log if folder deletion failed and user was still deleted.
            if not delete_success:
                logger.warning(
                    f"User object for '{username}' will be deleted, but their data folder deletion failed or had issues. Message: {delete_message}"
                )

            user_to_delete.delete()
            messages.success(request, f"用户 '{username}' 的账户已成功删除。")
    else:
        messages.warning(request, "删除操作应通过POST请求执行。")

    return redirect("converter:admin_user_management")


@login_required
@user_passes_test(is_superuser)
def admin_edit_user(request, user_id):
    user_to_edit = get_object_or_404(User, pk=user_id)

    if request.method == "POST":
        # 根据提交的表单类型分别处理
        if "change_info" in request.POST:
            user_form = AdminUserEditForm(request.POST, instance=user_to_edit)
            password_form = AdminSetPasswordForm(user_to_edit)  # 保持密码表单在上下文
            if user_form.is_valid():
                user_form.save()
                messages.success(
                    request, f"用户 '{user_to_edit.username}' 的信息已更新。"
                )
                return redirect("converter:admin_user_management")
        elif "set_password" in request.POST:
            password_form = AdminSetPasswordForm(user_to_edit, request.POST)
            user_form = AdminUserEditForm(
                instance=user_to_edit
            )  # 保持用户信息表单在上下文
            if password_form.is_valid():
                password_form.save()
                messages.success(
                    request, f"用户 '{user_to_edit.username}' 的密码已重置。"
                )
                return redirect("converter:admin_user_management")
        else:
            # 未知POST请求或缺少标识，可以简单地重新加载表单
            user_form = AdminUserEditForm(instance=user_to_edit)
            password_form = AdminSetPasswordForm(user_to_edit)
            messages.error(request, "无效的请求。")

    else:
        user_form = AdminUserEditForm(instance=user_to_edit)
        password_form = AdminSetPasswordForm(user_to_edit)

    return render(
        request,
        "custom_admin/user_edit_form.html",
        {
            "user_form": user_form,
            "password_form": password_form,
            "user_to_edit": user_to_edit,
        },
    )


def append_document(source_doc, target_doc):
    """Appends content of source_doc to target_doc."""
    for element in source_doc.element.body:
        target_doc.element.body.append(element)


# --- 新的独立视图函数 ---


@login_required
@require_POST
def file_to_pdf_view(request):
    start_time_view = time.perf_counter()  # ADDED: Start timer
    request_id = generate_request_id()
    logger.info(f"file_to_pdf_view: Received request. RequestID: {request_id}")

    username = request.user.username
    today_date_str = datetime.now().strftime("%Y%m%d")

    user_upload_dir, user_converted_dir = ensure_user_directories(
        username, today_date_str
    )
    if not user_upload_dir:
        logger.error(
            f"file_to_pdf_view: Failed to ensure user directories for {username} on {today_date_str}. RequestID: {request_id}"
        )
        # ADDED duration_seconds to error response
        return format_error_response(
            message="无法创建用户目录，请联系管理员。",
            merge_output=request.POST.get("merge_output", "false") == "true",
            request_id=request_id,
            duration_seconds=round(time.perf_counter() - start_time_view, 2),
        )

    # robust 文件上传处理逻辑，和 img_to_file_view 保持一致，保证 uploaded_files_info 一定被定义
    uploaded_files_info = []
    if request.FILES.getlist("images"):
        for uploaded_file_obj in request.FILES.getlist("images"):
            temp_input_path, original_filename, safe_filename = save_uploaded_file(
                uploaded_file_obj, user_upload_dir, request_id
            )
            if temp_input_path and safe_filename:
                uploaded_files_info.append(
                    {
                        "name": original_filename,
                        "path": temp_input_path,
                        "safe_original_filename": safe_filename,
                        "status": "uploaded",
                    }
                )
            else:
                uploaded_files_info.append(
                    {
                        "name": getattr(uploaded_file_obj, "name", "未知文件"),
                        "path": None,
                        "safe_original_filename": None,
                        "status": "error",
                    }
                )
    else:
        uploaded_files_info_from_frontend = request.POST.getlist(
            "uploaded_files_info[]"
        )
        if uploaded_files_info_from_frontend:
            try:
                parsed_files_info = [
                    json.loads(info) for info in uploaded_files_info_from_frontend
                ]
            except json.JSONDecodeError as e:
                logger.error(
                    f"file_to_pdf_view: JSONDecodeError parsing uploaded_files_info: {e}. RequestID: {request_id}"
                )
                return format_error_response(
                    message=f"解析文件信息时出错: {e}",
                    merge_output=request.POST.get("merge_output", "false") == "true",
                    request_id=request_id,
                    duration_seconds=round(time.perf_counter() - start_time_view, 2),
                )
            uploaded_files_info = parsed_files_info
        else:
            logger.warning(
                f"file_to_pdf_view: No uploaded_files_info provided. RequestID: {request_id}"
            )
            return format_error_response(
                message="没有提供文件信息。",
                merge_output=request.POST.get("merge_output", "false") == "true",
                request_id=request_id,
                duration_seconds=round(time.perf_counter() - start_time_view, 2),
            )

    # 其余参数和逻辑
    file_results = []  # INITIALIZED
    errors_view = []  # INITIALIZED
    temp_files_to_clean_view = []

    mode = request.POST.get("mode", "single")
    output_filename_base = request.POST.get(
        "output_filename_base", "converted_document"
    )
    merge_output_flag = request.POST.get("merge_output", "false") == "true"

    logger.info(
        f"file_to_pdf_view: Mode={mode}, MergeOutput={merge_output_flag}, OutputBase='{output_filename_base}'. RequestID: {request_id}"
    )

    conversion_func = None
    target_extension = None

    if mode == "docx_to_pdf_mode":
        conversion_method = request.POST.get("conversion_method_word", "libreoffice")
        logger.info(
            f"file_to_pdf_view: Word to PDF method selected: {conversion_method}. RequestID: {request_id}"
        )
        target_extension = ".pdf"
        if conversion_method == "docx2pdf" and DOCX2PDF_AVAILABLE_IN_VIEW:
            conversion_func = docx_to_pdf_converter_internal
        else:
            if conversion_method == "docx2pdf" and not DOCX2PDF_AVAILABLE_IN_VIEW:
                msg = "DOCX to PDF (docx2pdf)不可用，将尝试LibreOffice。如果问题持续，请联系管理员。"
                logger.warning(f"file_to_pdf_view: {msg} RequestID: {request_id}")
                errors_view.append(msg)
            conversion_func = convert_word_to_pdf
            logger.info(
                f"file_to_pdf_view: 使用 convert_word_to_pdf 通过 LibreOffice 处理 Word 转 PDF。RequestID: {request_id}"
            )
    elif mode == "excel_to_pdf_mode":
        conversion_func = convert_excel_to_pdf
        target_extension = ".pdf"
        logger.info(
            f"file_to_pdf_view: Excel to PDF mode selected. RequestID: {request_id}"
        )
    elif mode == "ppt_to_pdf_mode":
        conversion_func = ppt_pdf_converter.convert_pptx_to_pdf
        target_extension = ".pdf"
        logger.info(
            f"file_to_pdf_view: PPT to PDF mode selected. RequestID: {request_id}"
        )
    elif mode == "txt_to_pdf_mode":
        conversion_func = convert_txt_to_pdf
        target_extension = ".pdf"
        font_name = request.POST.get("font_name_txt", "SimSun")
        if conversion_func:
            original_conversion_func = conversion_func
            conversion_func = lambda input_path, output_path: original_conversion_func(
                input_path, output_path, font_name=font_name
            )
        logger.info(
            f"file_to_pdf_view: TXT to PDF mode selected with font: {font_name}. RequestID: {request_id}"
        )
    else:
        logger.error(
            f"file_to_pdf_view: Unknown or unsupported mode: {mode}. RequestID: {request_id}"
        )
        errors_view.append(f"未知的转换模式: {mode}")

    converted_pdf_paths_for_merge = []

    if not errors_view:
        for file_info in uploaded_files_info:
            original_filename = file_info.get("name", f"unknown_file_{request_id}")
            temp_input_path = file_info.get("path")

            if not temp_input_path or not os.path.exists(temp_input_path):
                logger.error(
                    f"file_to_pdf_view: Temporary input file not found or path is invalid for {original_filename} at {temp_input_path}. RequestID: {request_id}"
                )
                error_message = f"文件 '{original_filename}' 的临时路径无效或文件不存在，已跳过。"
                errors_view.append(error_message)
                file_results.append(
                    {
                        "original_name": original_filename,
                        "status": "error",
                        "message": error_message,
                    }
                )
                continue
            if not target_extension:
                logger.error(
                    f"file_to_pdf_view: Target extension not set for mode {mode}. File: {original_filename}. RequestID: {request_id}"
                )
                error_message = (
                    f"模式 {mode} 的目标文件类型未设置，无法处理 '{original_filename}'。"
                )
                errors_view.append(error_message)
                file_results.append(
                    {
                        "original_name": original_filename,
                        "status": "error",
                        "message": error_message,
                    }
                )
                continue

            safe_original_filename_base = Path(original_filename).stem
            safe_original_filename_base = re.sub(
                r"[^a-zA-Z0-9_\-\.]", "_", safe_original_filename_base
            )
            timestamp_suffix = datetime.now().strftime("%H%M%S%f")
            output_filename_display = f"{safe_original_filename_base}_{request_id}_{timestamp_suffix}{target_extension}"
            output_filepath_server = os.path.join(
                user_converted_dir, output_filename_display
            )

            if conversion_func:
                try:
                    logger.info(
                        f"file_to_pdf_view: Attempting conversion for {original_filename} to {output_filepath_server} using mode {mode}. RequestID: {request_id}"
                    )

                    conversion_result = conversion_func(
                        temp_input_path, output_filepath_server
                    )

                    if isinstance(conversion_result, tuple):
                        # Some converters return (success, output_path, extra_info)
                        if len(conversion_result) >= 2:
                            conversion_success = bool(conversion_result[0])
                            actual_output_path = conversion_result[1]
                        else:
                            conversion_success = bool(conversion_result[0])
                            actual_output_path = output_filepath_server
                    else:
                        conversion_success = conversion_result is None or bool(
                            conversion_result
                        )
                        actual_output_path = output_filepath_server

                    if not actual_output_path:
                        actual_output_path = output_filepath_server

                    if not conversion_success or not os.path.exists(actual_output_path):
                        logger.error(
                            f"file_to_pdf_view: Conversion did not produce expected file at {actual_output_path}. RequestID: {request_id}"
                        )
                        error_message = (
                            f"文件 '{original_filename}' 转换失败，未生成目标文件。"
                        )
                        errors_view.append(error_message)
                        file_results.append(
                            {
                                "original_name": original_filename,
                                "status": "error",
                                "message": error_message,
                            }
                        )
                        continue

                    logger.info(
                        f"file_to_pdf_view: Conversion successful for {original_filename}. Output: {actual_output_path}. RequestID: {request_id}"
                    )
                    file_results.append(
                        {
                            "original_name": original_filename,
                            "converted_name": os.path.basename(actual_output_path),
                            "download_url": reverse(
                                "converter:download_converted_file",
                                args=[
                                    username,
                                    today_date_str,
                                    os.path.basename(actual_output_path),
                                ],
                            ),
                            "status": "success",
                            "message": f"成功转换为 {target_extension.upper()} 文件。",
                        }
                    )
                    if merge_output_flag and target_extension == ".pdf":
                        converted_pdf_paths_for_merge.append(actual_output_path)
                except NotImplementedError as e_ni:
                    logger.error(
                        f"file_to_pdf_view: NotImplementedError for {original_filename}: {e_ni}. RequestID: {request_id}"
                    )
                    error_message = (
                        f"'{original_filename}' 的转换功能未实现或不可用: {e_ni}"
                    )
                    errors_view.append(error_message)
                    file_results.append(
                        {
                            "original_name": original_filename,
                            "status": "error",
                            "message": error_message,
                        }
                    )
                except Exception as e_conv:
                    logger.error(
                        f"file_to_pdf_view: Conversion error for {original_filename}: {e_conv}. Traceback: {traceback.format_exc()}. RequestID: {request_id}"
                    )
                    error_message = f"处理 '{original_filename}' 时出错: {e_conv}"
                    errors_view.append(error_message)
                    file_results.append(
                        {
                            "original_name": original_filename,
                            "status": "error",
                            "message": error_message,
                        }
                    )
            else:
                logger.error(
                    f"file_to_pdf_view: No conversion function resolved for mode {mode}, file {original_filename}. RequestID: {request_id}"
                )
                error_message = (
                    f"无法为 '{original_filename}' (模式: {mode}) 找到合适的转换器。"
                )
                errors_view.append(error_message)
                file_results.append(
                    {
                        "original_name": original_filename,
                        "status": "error",
                        "message": error_message,
                    }
                )

    if (
        merge_output_flag
        and mode
        in [
            "docx_to_pdf_mode",
            "excel_to_pdf_mode",
            "ppt_to_pdf_mode",
            "txt_to_pdf_mode",
        ]
        and converted_pdf_paths_for_merge
    ):
        if not PYPDF2_AVAILABLE:
            logger.warning(
                f"file_to_pdf_view: PyPDF2 not available, cannot merge PDFs. RequestID: {request_id}"
            )
            errors_view.append("PDF合并功能不可用 (缺少PyPDF2库)。单个文件已转换。")
        elif len(converted_pdf_paths_for_merge) == 1:
            # 只有一个文件，直接用它作为合并输出，不再复制
            single_pdf = converted_pdf_paths_for_merge[0]
            original_name = uploaded_files_info[0].get(
                "name", os.path.basename(single_pdf)
            )
            file_results = [
                {
                    "original_name": original_name,
                    "converted_name": os.path.basename(single_pdf),
                    "download_url": reverse(
                        "converter:download_converted_file",
                        args=[username, today_date_str, os.path.basename(single_pdf)],
                    ),
                    "status": "success",
                    "message": "仅有一个文件，已直接作为合并结果输出。",
                }
            ]
            # 不要清理 single_pdf（最终输出文件）
            # temp_files_to_clean_view.append(single_pdf)
        elif len(converted_pdf_paths_for_merge) > 1:
            merge_success = False
            merged_filename_display = f"{output_filename_base}_{request_id}_merged.pdf"
            merged_filepath_server = os.path.join(
                user_converted_dir, merged_filename_display
            )
            try:
                merger = PdfMerger()
                for pdf_path in converted_pdf_paths_for_merge:
                    if os.path.exists(pdf_path):
                        with open(pdf_path, "rb") as f_pdf:
                            reader = PdfReader(f_pdf)
                            if reader.is_encrypted:
                                try:
                                    reader.decrypt("")
                                except Exception as e_decrypt:
                                    logger.warning(
                                        f"Could not decrypt {pdf_path} with empty password: {e_decrypt}. Skipping file in merge."
                                    )
                                    errors_view.append(
                                        f"文件 {os.path.basename(pdf_path)} 已加密且无法解密，已在合并中跳过。"
                                    )
                                    continue
                            merger.append(reader)
                    else:
                        logger.warning(
                            f"file_to_pdf_view: File {pdf_path} not found for merging. RequestID: {request_id}"
                        )

                if merger.pages:
                    merger.write(merged_filepath_server)
                    merger.close()
                    merge_success = True
                    logger.info(
                        f"file_to_pdf_view: Successfully merged {len(converted_pdf_paths_for_merge)} PDFs into {merged_filepath_server}. RequestID: {request_id}"
                    )

                    # Create a descriptive name showing source files
                    source_files = [
                        file_info.get("name", "unknown")
                        for file_info in uploaded_files_info[:3]
                    ]  # Show up to 3 file names
                    if len(uploaded_files_info) > 3:
                        source_files.append(f"等{len(uploaded_files_info)}个文件")
                    source_names = "、".join(source_files)
                    original_name_display = f"合并PDF (来自: {source_names})"

                    file_results = [
                        {
                            "original_name": original_name_display,
                            "converted_name": merged_filename_display,
                            "download_url": reverse(
                                "converter:download_converted_file",
                                args=[
                                    username,
                                    today_date_str,
                                    merged_filename_display,
                                ],
                            ),
                            "status": "success",
                            "message": f"成功合并 {len(converted_pdf_paths_for_merge)} 个文件为一个PDF。",
                        }
                    ]
                    # 不要清理 merged_filepath_server（最终输出文件）
                    # temp_files_to_clean_view.append(merged_filepath_server)
                else:
                    logger.warning(
                        f"file_to_pdf_view: No pages were added to the merger. Merged file not created. RequestID: {request_id}"
                    )
                    if not errors_view:
                        errors_view.append(
                            "未能合并PDF文件（没有内容可合并）。单个文件可能已转换。"
                        )
            except Exception as e_merge:
                logger.error(
                    f"file_to_pdf_view: Error merging PDFs: {e_merge}. Traceback: {traceback.format_exc()}. RequestID: {request_id}"
                )
                errors_view.append(f"合并PDF时出错: {e_merge}")
                merge_success = False

            if merge_success:
                for pdf_path in converted_pdf_paths_for_merge:
                    temp_files_to_clean_view.append(pdf_path)

    for file_info in uploaded_files_info:
        if file_info.get("path") and os.path.exists(file_info.get("path")):
            temp_files_to_clean_view.append(file_info.get("path"))

    if temp_files_to_clean_view:
        cleanup_temp_files(temp_files_to_clean_view, request_id)

    if not file_results and not errors_view:
        logger.warning(
            f"file_to_pdf_view: No results and no errors. This might indicate an issue. RequestID: {request_id}"
        )
        errors_view.append("没有文件被成功处理，也没有明确的错误信息。")

    end_time_view = time.perf_counter()  # ADDED: End timer
    duration_seconds_view = round(
        end_time_view - start_time_view, 2
    )  # ADDED: Calculate duration

    error_summary = "; ".join(errors_view) if errors_view else None
    if errors_view:
        logger.error(
            f"file_to_pdf_view: Processing finished with errors. Duration: {duration_seconds_view}s. Errors: {errors_view}. RequestID: {request_id}"
        )

    logger.info(
        f"file_to_pdf_view: Processing complete. Duration: {duration_seconds_view}s. RequestID: {request_id}"
    )
    return format_json_response(
        results=file_results,
        merge_output=merge_output_flag,
        request_id=request_id,
        duration_seconds=duration_seconds_view,
        error_message=None if file_results else error_summary,
    )


@login_required
@require_POST
def img_to_file_view(request):
    start_time_view = time.perf_counter()  # ADDED: Start timer
    request_id = generate_request_id()
    logger.info(f"img_to_file_view: Received request. RequestID: {request_id}")

    username = request.user.username
    today_date_str = datetime.now().strftime("%Y%m%d")
    errors_view = []  # INITIALIZED errors_view

    user_upload_dir, user_converted_dir = "", ""
    try:
        user_upload_dir, user_converted_dir = ensure_user_directories(
            username, today_date_str
        )
    except Exception as e:
        logger.critical(
            f"img_to_file_view: Failed to create user directories for {username}. Error: {e}. RequestID: {request_id}",
            exc_info=True,
        )
        merge_output_for_error = (
            request.POST.get("merge_output", "false").lower() == "true"
        )
        return format_error_response(
            message="服务器错误：无法创建用户目录。",
            merge_output=merge_output_for_error,
            request_id=request_id,
        )

    parsed_params = {}
    try:
        parsed_params = parse_conversion_request_params(request.POST, request_id)
    except Exception as e_parse:
        logger.error(
            f"img_to_file_view: Error parsing request parameters: {e_parse}. RequestID: {request_id}",
            exc_info=True,
        )
        merge_output_for_error = (
            request.POST.get("merge_output", "false").lower() == "true"
        )
        return format_error_response(
            message="请求参数错误。",
            merge_output=merge_output_for_error,
            request_id=request_id,
        )

    merge_output = parsed_params["merge_output"]
    output_format = parsed_params["output_format"]

    processed_files_final = []
    temp_files_to_delete_final = []

    uploaded_files_info = []
    if request.FILES.getlist("images"):
        for uploaded_file_obj in request.FILES.getlist("images"):
            logger.info(
                f"Processing upload: {getattr(uploaded_file_obj, 'name', 'no name')}"
            )
            temp_input_path, original_filename, safe_filename = save_uploaded_file(
                uploaded_file_obj, user_upload_dir, request_id
            )
            logger.info(f"Saved to: {temp_input_path}")
            if temp_input_path and safe_filename:
                uploaded_files_info.append(
                    {
                        "name": original_filename,
                        "path": temp_input_path,
                        "safe_original_filename": safe_filename,
                        "status": "uploaded",
                    }
                )
            else:
                uploaded_files_info.append(
                    {
                        "name": getattr(uploaded_file_obj, "name", "未知文件"),
                        "path": None,
                        "safe_original_filename": None,
                        "status": "error",
                    }
                )
    else:
        uploaded_files_info_from_frontend = request.POST.getlist(
            "uploaded_files_info[]"
        )
        if uploaded_files_info_from_frontend:
            try:
                parsed_files_info = [
                    json.loads(info) for info in uploaded_files_info_from_frontend
                ]
            except json.JSONDecodeError as e:
                logger.error(
                    f"file_to_pdf_view: JSONDecodeError parsing uploaded_files_info: {e}. RequestID: {request_id}"
                )
                return format_error_response(
                    message=f"解析文件信息时出错: {e}",
                    merge_output=request.POST.get("merge_output", "false") == "true",
                    request_id=request_id,
                    duration_seconds=round(time.perf_counter() - start_time_view, 2),
                )
            uploaded_files_info = parsed_files_info
        else:
            logger.warning(
                f"file_to_pdf_view: No uploaded_files_info provided. RequestID: {request_id}"
            )
            return format_error_response(
                message="没有提供文件信息。",
                merge_output=request.POST.get("merge_output", "false") == "true",
                request_id=request_id,
                duration_seconds=round(time.perf_counter() - start_time_view, 2),
            )

    if not uploaded_files_info:
        logger.error(
            f"img_to_file_view: All file uploads failed or no files were valid after saving. RequestID: {request_id}"
        )
        if not processed_files_final:
            processed_files_final.append(
                {
                    "original_name": "File Upload",
                    "status": "error",
                    "message": "所有文件上传失败或未能保存。",
                }
            )
        return format_json_response(
            results=processed_files_final,
            merge_output=merge_output,
            request_id=request_id,
        )

    # 直接插图逻辑提前，确保优先分支
    direct_image_to_ppt = (
        request.POST.get("direct_image_to_ppt", "false").lower() == "true"
    )
    if output_format == "pptx" and direct_image_to_ppt:
        image_paths = [
            item["path"] for item in uploaded_files_info if os.path.exists(item["path"])
        ]
        if not image_paths:
            return format_error_response(
                message="没有可用的图片文件用于PPTX生成。",
                merge_output=merge_output,
                request_id=request_id,
            )
        pptx_filename = f"images_{request_id}.pptx"
        pptx_path = os.path.join(user_converted_dir, pptx_filename)
        try:
            copy_images_to_pptx(image_paths, pptx_path)
            processed_files_final = [
                {
                    "original_name": f"批量图片转PPTX ({len(image_paths)}张)",
                    "converted_name": pptx_filename,
                    "download_url": reverse(
                        "converter:download_converted_file",
                        args=[request.user.username, today_date_str, pptx_filename],
                    ),
                    "status": "success",
                    "message": "图片已直接插入PPT。",
                }
            ]
            cleanup_temp_files(image_paths, request_id)
            return format_json_response(
                results=processed_files_final,
                merge_output=merge_output,
                request_id=request_id,
            )
        except Exception as e:
            logger.error(f"图片直接插入PPTX失败: {e}", exc_info=True)
            return format_error_response(
                message=f"图片直接插入PPTX失败: {e}",
                merge_output=merge_output,
                request_id=request_id,
            )

    # --- Core Img to File conversion logic (formerly _handle_img_to_file) ---
    img_script_results, script_created_files = process_images_to_files(
        uploaded_files_info, user_converted_dir, request_id, output_format
    )

    for item in uploaded_files_info:
        if item.get("path") and os.path.exists(item["path"]):
            temp_files_to_delete_final.append(item["path"])

    if not script_created_files and not img_script_results:
        logger.error(
            f"img_to_file_view: pic_file_converter script (process_images_to_files) provided no output. RequestID: {request_id}"
        )
        if not img_script_results:  # Populate errors if script returned nothing
            for up_file in uploaded_files_info:
                processed_files_final.append(
                    {
                        "original_name": up_file["name"],
                        "status": "error",
                        "message": "图像处理脚本未能生成任何输出。",
                    }
                )
        else:  # Script might have returned error messages within img_script_results
            processed_files_final.extend(img_script_results)

        cleanup_temp_files(list(set(temp_files_to_delete_final)), request_id)
        return format_json_response(
            results=processed_files_final,
            merge_output=merge_output,
            request_id=request_id,
        )

    if merge_output:
        if not script_created_files:
            logger.warning(
                f"img_to_file_view: Merge requested, but no files from script to merge. Script results: {img_script_results}. RequestID: {request_id}"
            )
            processed_files_final.extend(
                img_script_results
                or [
                    {
                        "original_name": "图像合并操作",
                        "status": "error",
                        "message": "没有从图像生成可合并的文档。",
                    }
                ]
            )
            temp_files_to_delete_final.extend(
                [
                    item["path"]
                    for item in (script_created_files or [])
                    if isinstance(item, dict) and "path" in item
                ]
            )
        else:
            merged_base_name = f"merged_images_{request_id}"
            final_merged_docx_filename = f"{merged_base_name}.docx"
            final_merged_docx_path = os.path.join(
                user_converted_dir, final_merged_docx_filename
            )
            try:
                # 合并前检查所有待合并文件是否存在
                missing_files = [
                    f["path"]
                    for f in script_created_files
                    if not os.path.exists(f["path"])
                ]
                if missing_files:
                    logger.error(f"以下中间DOCX文件不存在，无法合并: {missing_files}")
                    raise FileNotFoundError(f"中间DOCX文件缺失: {missing_files}")
                master_doc = Document(script_created_files[0]["path"])
                for idx, doc_info in enumerate(script_created_files[1:], start=1):
                    sub_doc = Document(doc_info["path"])
                    append_document(sub_doc, master_doc)
                    # 只在不是最后一个子文档后插入分页符
                    if idx < len(script_created_files) - 1:
                        para = master_doc.add_paragraph()
                        para.add_run().add_break(break_type=WD_BREAK.PAGE)
                master_doc.save(final_merged_docx_path)
                logger.info(
                    f"img_to_file_view: Merged {len(script_created_files)} DOCX files to {final_merged_docx_path}. RequestID: {request_id}"
                )
                temp_files_to_delete_final.extend(
                    [
                        item["path"]
                        for item in (script_created_files or [])
                        if isinstance(item, dict) and "path" in item
                    ]
                )

                if output_format == "docx":
                    # Create a descriptive name showing source files
                    source_files = [
                        item["name"] for item in uploaded_files_info[:3]
                    ]  # Show up to 3 file names
                    if len(uploaded_files_info) > 3:
                        source_files.append(f"等{len(uploaded_files_info)}个文件")
                    source_names = "、".join(source_files)
                    original_name_display = f"合并文档 (来自: {source_names})"

                    processed_files_final = [
                        {
                            "original_name": original_name_display,
                            "converted_name": final_merged_docx_filename,
                            "download_url": reverse(
                                "converter:download_converted_file",
                                args=[
                                    request.user.username,
                                    today_date_str,
                                    final_merged_docx_filename,
                                ],
                            ),
                            "status": "success",
                            "message": "图像已成功合并为Word文档。",
                        }
                    ]
                elif output_format == "pdf":
                    final_merged_pdf_filename = f"{merged_base_name}.pdf"
                    final_merged_pdf_path = os.path.join(
                        user_converted_dir, final_merged_pdf_filename
                    )

                    pdf_success, pdf_path_or_msg, _ = convert_word_to_pdf(
                        final_merged_docx_path, final_merged_pdf_path
                    )
                    if pdf_success and os.path.exists(final_merged_pdf_path):
                        # Create a descriptive name showing source files
                        source_files = [
                            item["name"] for item in uploaded_files_info[:3]
                        ]  # Show up to 3 file names
                        if len(uploaded_files_info) > 3:
                            source_files.append(f"等{len(uploaded_files_info)}个文件")
                        source_names = "、".join(source_files)
                        original_name_display = f"合并PDF (来自: {source_names})"

                        processed_files_final = [
                            {
                                "original_name": original_name_display,
                                "converted_name": final_merged_pdf_filename,
                                "download_url": reverse(
                                    "converter:download_converted_file",
                                    args=[
                                        request.user.username,
                                        today_date_str,
                                        final_merged_pdf_filename,
                                    ],
                                ),
                                "status": "success",
                                "message": "图像成功合并到Word并转换为PDF。",
                            }
                        ]
                        temp_files_to_delete_final.append(final_merged_docx_path)
                    else:
                        processed_files_final = [
                            {
                                "original_name": "图像合并与PDF转换",
                                "status": "error",
                                "message": pdf_path_or_msg
                                or "无法将合并的Word文档转换为PDF。",
                            }
                        ]
                        # MODIFIED: Do not delete intermediate merged DOCX if output is PPTX for debugging
                        # Temporarily disable cleanup to inspect the merged DOCX file
                        # if output_format != 'pptx':                            #     temp_files_to_delete_final.append(final_merged_docx_path)
                elif (
                    output_format == "pptx"
                ):  # New: Handle PPTX output for merged files
                    final_merged_pptx_filename = f"{merged_base_name}.pptx"
                    final_merged_pptx_path = os.path.join(
                        user_converted_dir, final_merged_pptx_filename
                    )
                    pptx_success, pptx_path_or_msg, _ = (
                        convert_docx_to_pptx_libreoffice(
                            final_merged_docx_path,
                            user_converted_dir,
                            skip_default_content=True,
                        )
                    )

                    if (
                        pptx_success
                        and pptx_path_or_msg
                        and os.path.exists(pptx_path_or_msg)
                    ):
                        # pptx_path_or_msg from libreoffice converter is the actual path of the created file (e.g., user_converted_dir/merged_images_requestid.pptx)
                        # We need to rename it to final_merged_pptx_path if it's different (it should be if libreoffice names it based on docx stem)
                        if pptx_path_or_msg != final_merged_pptx_path:
                            if os.path.exists(final_merged_pptx_path):
                                os.remove(
                                    final_merged_pptx_path
                                )  # Remove if somehow exists
                            shutil.move(pptx_path_or_msg, final_merged_pptx_path)

                        # Create a descriptive name showing source files
                        source_files = [
                            item["name"] for item in uploaded_files_info[:3]
                        ]  # Show up to 3 file names
                        if len(uploaded_files_info) > 3:
                            source_files.append(f"等{len(uploaded_files_info)}个文件")
                        source_names = "、".join(source_files)
                        original_name_display = f"合并PPTX (来自: {source_names})"

                        processed_files_final = [
                            {
                                "original_name": original_name_display,
                                "converted_name": final_merged_pptx_filename,
                                "download_url": reverse(
                                    "converter:download_converted_file",
                                    args=[
                                        request.user.username,
                                        today_date_str,
                                        final_merged_pptx_filename,
                                    ],
                                ),
                                "status": "success",
                                "message": "图像成功合并到Word并转换为PPTX。",
                            }
                        ]
                        # MODIFIED: Do not delete intermediate merged DOCX if output is PPTX for debugging
                        # Temporarily disable cleanup to inspect the merged DOCX file
                        # if output_format != 'pptx':
                        #     temp_files_to_delete_final.append(final_merged_docx_path)
                    else:
                        processed_files_final = [
                            {
                                "original_name": "图像合并与PPTX转换",
                                "status": "error",
                                "message": pptx_path_or_msg
                                or "无法将合并的Word文档转换为PPTX。",
                            }
                        ]
                        # MODIFIED: Do not delete intermediate merged DOCX if output is PPTX for debugging, and it was an error with PPTX conversion
                        # However, if the error is about PPTX conversion, the DOCX might be useful.
                        # Let's keep it for now if output_format == 'pptx'. If it's another format, it should be deleted.
                        # if output_format != 'pptx':                            #     temp_files_to_delete_final.append(final_merged_docx_path)
                else:
                    processed_files_final.extend(
                        img_script_results
                        or [
                            {
                                "original_name": "图像合并操作",
                                "status": "info",
                                "message": "请求合并，但没有生成可合并的图像文档。",
                            }
                        ]
                    )
            except Exception as e_img_merge:
                logger.error(
                    f"img_to_file_view: 合并DOCX文件时出错: {e_img_merge}. RequestID: {request_id}",
                    exc_info=True,
                )
                processed_files_final.append(
                    {
                        "original_name": "图像合并操作",
                        "status": "error",
                        "message": f"图像文档合并过程中出错: {str(e_img_merge)}",
                    }
                )
                if os.path.exists(final_merged_docx_path):
                    temp_files_to_delete_final.append(final_merged_docx_path)
                temp_files_to_delete_final.extend(
                    [
                        item["path"]
                        for item in (script_created_files or [])
                        if isinstance(item, dict) and "path" in item
                    ]
                )
    else:
        temp_individual_results = []
        if not img_script_results:
            logger.warning(
                f"img_to_file_view: 非合并模式，但 process_images_to_files 未返回结果。RequestID: {request_id}"
            )
            for (
                up_file
            ) in uploaded_files_info:  # Create error entries for each uploaded file
                temp_individual_results.append(
                    {
                        "original_name": up_file["name"],
                        "status": "error",
                        "message": "图像处理脚本未能为此文件生成输出。",
                    }
                )
            temp_files_to_delete_final.extend(
                [
                    item["path"]
                    for item in (script_created_files or [])
                    if isinstance(item, dict) and "path" in item
                ]
            )
        else:
            for res_info in img_script_results:
                original_img_name = res_info.get("original_name")
                intermediate_docx_name = res_info.get("converted_name")
                intermediate_docx_full_path = None
                if intermediate_docx_name:  # Map to full path
                    for scf_item in script_created_files:  # scf_item is a dict
                        if (
                            isinstance(scf_item, dict)
                            and os.path.basename(scf_item.get("path", ""))
                            == intermediate_docx_name
                        ):
                            intermediate_docx_full_path = scf_item["path"]
                            break

                if (
                    res_info.get("status") == "error"
                    or not intermediate_docx_full_path
                    or not os.path.exists(intermediate_docx_full_path)
                ):
                    temp_individual_results.append(
                        res_info
                    )  # Pass through script's error or if file is missing
                    if intermediate_docx_full_path and os.path.exists(
                        intermediate_docx_full_path
                    ):  # Cleanup if exists but was error
                        temp_files_to_delete_final.append(intermediate_docx_full_path)
                    elif intermediate_docx_full_path and not os.path.exists(
                        intermediate_docx_full_path
                    ):
                        logger.warning(
                            f"img_to_file_view: 脚本为 {original_img_name} 生成的 {intermediate_docx_name} 不存在于 {intermediate_docx_full_path}. RequestID: {request_id}"
                        )
                    continue

                if output_format == "docx":
                    res_info["download_url"] = reverse(
                        "converter:download_converted_file",
                        args=[
                            request.user.username,
                            today_date_str,
                            intermediate_docx_name,
                        ],
                    )
                    temp_individual_results.append(res_info)
                elif output_format == "pdf":
                    pdf_base = os.path.splitext(intermediate_docx_name)[0]
                    final_pdf_name = f"{pdf_base}.pdf"
                    final_pdf_full_path = os.path.join(
                        user_converted_dir, final_pdf_name
                    )
                    pdf_succ, pdf_msg, _ = convert_word_to_pdf(
                        intermediate_docx_full_path, final_pdf_full_path
                    )
                    if pdf_succ and os.path.exists(final_pdf_full_path):
                        temp_individual_results.append(
                            {
                                "original_name": original_img_name,
                                "converted_name": final_pdf_name,
                                "download_url": reverse(
                                    "converter:download_converted_file",
                                    args=[
                                        request.user.username,
                                        today_date_str,
                                        final_pdf_name,
                                    ],
                                ),
                                "status": "success",
                                "message": res_info.get("message", "图像已转为PDF。")
                                + " (经Word)",
                            }
                        )
                        temp_files_to_delete_final.append(intermediate_docx_full_path)
                    else:
                        temp_individual_results.append(
                            {
                                "original_name": original_img_name,
                                "converted_name": intermediate_docx_name,
                                "download_url": reverse(
                                    "converter:download_converted_file",
                                    args=[
                                        request.user.username,
                                        today_date_str,
                                        intermediate_docx_name,
                                    ],
                                ),
                                "status": "error",
                                "message": pdf_msg or "图像生成的Word转PDF失败。",
                            }
                        )
                elif (
                    output_format == "pptx"
                ):  # New: Handle PPTX output for individual files
                    pptx_base_name_no_ext = os.path.splitext(intermediate_docx_name)[
                        0
                    ]  # original_img_name_tempScriptOutput_requestid
                    # We want the final name to be like: original_img_name_requestid.pptx
                    # The intermediate_docx_name is like: original_img_name_tempScriptOutput_requestid.docx
                    # So, pptx_base_name_no_ext is original_img_name_tempScriptOutput_requestid
                    # Let's try to reconstruct a cleaner name if possible, or use a unique one.
                    # For consistency, use original_img_name and request_id for the final pptx name
                    final_pptx_name = (
                        f"{os.path.splitext(original_img_name)[0]}_{request_id}.pptx"
                    )
                    final_pptx_full_path = os.path.join(
                        user_converted_dir, final_pptx_name
                    )
                    pptx_succ, actual_libre_pptx_path, pptx_msg = (
                        convert_docx_to_pptx_libreoffice(
                            intermediate_docx_full_path,
                            user_converted_dir,
                            skip_default_content=True,
                        )
                    )

                    if (
                        pptx_succ
                        and actual_libre_pptx_path
                        and os.path.exists(actual_libre_pptx_path)
                    ):
                        # actual_libre_pptx_path is based on intermediate_docx_full_path's stem, e.g. user_converted_dir/original_img_name_tempScriptOutput_requestid.pptx
                        # We need to rename it to final_pptx_full_path
                        if actual_libre_pptx_path != final_pptx_full_path:
                            if os.path.exists(final_pptx_full_path):
                                os.remove(final_pptx_full_path)
                            shutil.move(actual_libre_pptx_path, final_pptx_full_path)

                        temp_individual_results.append(
                            {
                                "original_name": original_img_name,
                                "converted_name": final_pptx_name,
                                "download_url": reverse(
                                    "converter:download_converted_file",
                                    args=[
                                        request.user.username,
                                        today_date_str,
                                        final_pptx_name,
                                    ],
                                ),
                                "status": "success",
                                "message": res_info.get("message", "图像已转为PPTX。")
                                + " (经Word)",
                            }
                        )
                        # MODIFIED: Do not delete intermediate merged DOCX if output is PPTX for debugging
                        # Temporarily disable cleanup to inspect the merged DOCX file
                        # if output_format != 'pptx':
                        #     temp_files_to_delete_final.append(final_merged_docx_path)
                    else:
                        temp_individual_results.append(
                            {
                                "original_name": original_img_name,
                                "converted_name": intermediate_docx_name,
                                "download_url": reverse(
                                    "converter:download_converted_file",
                                    args=[
                                        request.user.username,
                                        today_date_str,
                                        intermediate_docx_name,
                                    ],
                                ),
                                "status": "error",
                                "message": pptx_msg or "图像生成的Word转PPTX失败。",
                            }
                        )
                        # MODIFIED: Do not delete intermediate DOCX if output is PPTX for debugging, and it was an error with PPTX conversion
                        # However, if the error is about PPTX conversion, the DOCX might be useful.
                        # Let's keep it for now if output_format == 'pptx'. If it's another format, it should be deleted.
                        # if output_format != 'pptx':                            #     temp_files_to_delete_final.append(final_merged_docx_path)
                else:
                    logger.warning(
                        f"img_to_file_view: 不支持的输出格式 '{output_format}' 用于单个图像处理. RequestID: {request_id}"
                    )
                    res_info["status"] = "error"
                    res_info["message"] = f"图像转换不支持输出格式 '{output_format}'。"
                    temp_individual_results.append(res_info)

        processed_files_final = temp_individual_results
        final_product_names = [
            f.get("converted_name")
            for f in processed_files_final
            if f.get("status") == "success"
        ]
        for (
            scf_item
        ) in (
            script_created_files
        ):  # General cleanup of unmerged/unused script files; scf_item is a dict
            if isinstance(scf_item, dict) and "path" in scf_item:
                scf_path = scf_item["path"]
                if os.path.basename(scf_path) not in final_product_names:
                    if scf_path not in temp_files_to_delete_final:
                        temp_files_to_delete_final.append(scf_path)

    cleanup_temp_files(list(set(temp_files_to_delete_final)), request_id)
    end_time_view = time.perf_counter()  # ADDED: End timer
    duration_seconds_view = round(
        end_time_view - start_time_view, 2
    )  # ADDED: Calculate duration

    if errors_view:  # Check errors_view here
        logger.error(
            f"img_to_file_view: Processing finished with errors. Duration: {duration_seconds_view}s. Errors: {errors_view}. RequestID: {request_id}"
        )
        # Ensure format_error_response is called with merge_output if it expects it
        # Looking at its definition, it doesn't strictly require merge_output, but other calls include it.
        # For consistency, let's try to determine merge_output if possible, or pass a default.
        merge_output_for_error_response = parsed_params.get(
            "merge_output", False
        )  # Get from parsed_params if available
        return format_error_response(
            message="; ".join(errors_view),
            request_id=request_id,
            duration_seconds=duration_seconds_view,
            merge_output=merge_output_for_error_response,
        )

    # Use processed_files_final for results
    final_result_payload = {
        "results": processed_files_final,  # CHANGED from file_results
        "request_id": request_id,
        "merge_output": parsed_params.get(
            "merge_output", False
        ),  # Get from parsed_params
        "duration_seconds": duration_seconds_view,
    }
    logger.info(
        f"img_to_file_view: Processing complete. Duration: {duration_seconds_view}s. Results: {len(processed_files_final)} files. RequestID: {request_id}"
    )  # CHANGED from file_results
    return format_json_response(
        results=final_result_payload["results"],
        merge_output=final_result_payload["merge_output"],
        request_id=final_result_payload["request_id"],
        duration_seconds=final_result_payload["duration_seconds"],
    )


@login_required
@require_POST
def pdf_to_file_view(request):
    start_time_view = time.perf_counter()  # ADDED: Start timer
    request_id = generate_request_id()
    logger.info(f"pdf_to_file_view: Received request. RequestID: {request_id}")

    username = request.user.username
    today_date_str = datetime.now().strftime("%Y%m%d")

    user_upload_dir, user_converted_dir = "", ""
    try:
        user_upload_dir, user_converted_dir = ensure_user_directories(
            username, today_date_str
        )
    except Exception as e:
        logger.critical(
            f"pdf_to_file_view: Failed to create user directories for {username}. Error: {e}. RequestID: {request_id}",
            exc_info=True,
        )
        merge_output_for_error = (
            request.POST.get("merge_output", "false").lower() == "true"
        )
        return format_error_response(
            message="服务器错误：无法创建用户目录。",
            merge_output=merge_output_for_error,
            request_id=request_id,
        )

    parsed_params = {}
    try:
        parsed_params = parse_conversion_request_params(request.POST, request_id)
    except Exception as e_parse:
        logger.error(
            f"pdf_to_file_view: Error parsing request parameters: {e_parse}. RequestID: {request_id}",
            exc_info=True,
        )
        merge_output_for_error = (
            request.POST.get("merge_output", "false").lower() == "true"
        )
        return format_error_response(
            message="请求参数错误。",
            merge_output=merge_output_for_error,
            request_id=request_id,
        )

    merge_output = parsed_params["merge_output"]
    sub_tab = parsed_params["sub_tab"]

    processed_files_final = []
    temp_files_to_delete_final = []

    uploaded_files_info = []
    if request.FILES.getlist("images"):
        for uploaded_file_obj in request.FILES.getlist("images"):
            temp_input_path, original_filename, safe_filename = save_uploaded_file(
                uploaded_file_obj, user_upload_dir, request_id
            )
            if temp_input_path and safe_filename:
                uploaded_files_info.append(
                    {
                        "name": original_filename,
                        "path": temp_input_path,
                        "safe_original_filename": safe_filename,
                        "status": "uploaded",
                    }
                )
            else:
                uploaded_files_info.append(
                    {
                        "name": getattr(uploaded_file_obj, "name", "未知文件"),
                        "path": None,
                        "safe_original_filename": None,
                        "status": "error",
                    }
                )
    else:
        uploaded_files_info_from_frontend = request.POST.getlist(
            "uploaded_files_info[]"
        )
        if uploaded_files_info_from_frontend:
            try:
                parsed_files_info = [
                    json.loads(info) for info in uploaded_files_info_from_frontend
                ]
            except json.JSONDecodeError as e:
                logger.error(
                    f"file_to_pdf_view: JSONDecodeError parsing uploaded_files_info: {e}. RequestID: {request_id}"
                )
                return format_error_response(
                    message=f"解析文件信息时出错: {e}",
                    merge_output=request.POST.get("merge_output", "false") == "true",
                    request_id=request_id,
                    duration_seconds=round(time.perf_counter() - start_time_view, 2),
                )
            uploaded_files_info = parsed_files_info
        else:
            logger.warning(
                f"file_to_pdf_view: No uploaded_files_info provided. RequestID: {request_id}"
            )
            return format_error_response(
                message="没有提供文件信息。",
                merge_output=request.POST.get("merge_output", "false") == "true",
                request_id=request_id,
                duration_seconds=round(time.perf_counter() - start_time_view, 2),
            )

    if not uploaded_files_info:
        logger.error(
            f"pdf_to_file_view: All PDF uploads failed. RequestID: {request_id}"
        )
        if not processed_files_final:
            processed_files_final.append(
                {
                    "original_name": "PDF File Upload",
                    "status": "error",
                    "message": "所有PDF文件上传失败或未能保存。",
                }
            )
        return format_json_response(
            results=processed_files_final,
            merge_output=merge_output,
            request_id=request_id,
        )

    # --- Core PDF to File conversion logic (formerly _handle_pdf_to_file) ---
    temp_individual_converted_outputs = []
    for up_file_info in uploaded_files_info:
        original_name = up_file_info["name"]
        source_file_path = up_file_info["path"]
        safe_original_filename = up_file_info["safe_original_filename"]
        base_name_no_ext = os.path.splitext(safe_original_filename)[0]

        converted_filename = None
        success_conv = False
        conversion_message_conv = "不支持的转换或发生错误。"
        actual_output_path_from_converter = None
        intended_final_output_path = None

        try:
            if not original_name.lower().endswith(".pdf"):
                error_message = f"文件类型不匹配: {original_name} (应为PDF)。"
                logger.warning(f"pdf_to_file_view: {error_message} RID: {request_id}")
                processed_files_final.append(
                    {
                        "original_name": original_name,
                        "status": "error",
                        "message": error_message,
                    }
                )
                if source_file_path and os.path.exists(source_file_path):
                    temp_files_to_delete_final.append(source_file_path)
                continue

            logger.info(
                f"pdf_to_file_view: Processing {original_name} for {sub_tab}. Input: {source_file_path}. RID: {request_id}"
            )

            output_extension = ""
            current_mode = None
            if sub_tab == "pdfToWord":
                output_extension = ".docx"
                current_mode = parsed_params["pdf_to_word_mode"]
            elif sub_tab == "pdfToExcel":
                output_extension = ".xlsx"
                current_mode = parsed_params["pdf_to_excel_mode"]
            elif sub_tab == "pdfToPpt":
                output_extension = ".pptx"
                current_mode = parsed_params["pdf_to_ppt_mode"]
            elif sub_tab == "pdfToTxt":
                output_extension = ".txt"
                current_mode = parsed_params["pdf_to_txt_mode"]

            if not output_extension:
                conversion_message_conv = f"不支持的PDF转换子类型: {sub_tab}。"
                logger.warning(
                    f"pdf_to_file_view: {conversion_message_conv} for {original_name}. RID: {request_id}"
                )
            else:
                converted_filename = (
                    f"{base_name_no_ext}_{request_id}{output_extension}"
                )
                intended_final_output_path = os.path.join(
                    user_converted_dir, converted_filename
                )

                if sub_tab == "pdfToWord":
                    (
                        success_conv,
                        actual_output_path_from_converter,
                        conversion_message_conv,
                    ) = convert_pdf_to_word(
                        source_file_path, intended_final_output_path, mode=current_mode
                    )
                elif sub_tab == "pdfToExcel":
                    (
                        success_conv,
                        actual_output_path_from_converter,
                        conversion_message_conv,
                    ) = convert_pdf_to_excel(
                        source_file_path, intended_final_output_path, mode=current_mode
                    )
                elif sub_tab == "pdfToPpt":
                    # Pass user_converted_dir as the output folder, and base_name_no_ext as desired_filename_base
                    # The converter will create a file like base_name_no_ext.pptx or base_name_no_ext_converted.pptx in user_converted_dir
                    (
                        success_conv,
                        actual_output_path_from_converter,
                        conversion_message_conv,
                    ) = convert_pdf_to_ppt(
                        source_file_path,
                        user_converted_dir,  # Pass the directory here
                        mode=current_mode,
                        desired_filename_base=base_name_no_ext,  # Pass the base name for the converter to use
                    )
                    # intended_final_output_path is already defined correctly with request_id
                elif sub_tab == "pdfToTxt":
                    (
                        success_conv,
                        actual_output_path_from_converter,
                        conversion_message_conv,
                    ) = convert_pdf_to_txt(
                        source_file_path, intended_final_output_path, mode=current_mode
                    )

            if (
                success_conv
                and actual_output_path_from_converter
                and os.path.exists(actual_output_path_from_converter)
            ):
                # Ensure actual_output_path_from_converter is a file before proceeding
                if not os.path.isfile(actual_output_path_from_converter):
                    success_conv = False
                    conversion_message_conv = f"转换器未返回有效的文件路径: {actual_output_path_from_converter}"
                    logger.error(
                        f"pdf_to_file_view/{sub_tab}: Converter for '{original_name}' returned a non-file path: {actual_output_path_from_converter}. RID: {request_id}"
                    )
                elif actual_output_path_from_converter != intended_final_output_path:
                    if os.path.exists(intended_final_output_path):
                        try:
                            if os.path.isdir(intended_final_output_path):
                                shutil.rmtree(
                                    intended_final_output_path
                                )  # Remove directory if it exists from previous error
                            else:
                                os.remove(
                                    intended_final_output_path
                                )  # Remove file if it exists
                        except Exception as e_remove_existing:
                            logger.warning(
                                f"Error removing existing target {intended_final_output_path} before move: {e_remove_existing}. RID: {request_id}"
                            )
                    try:
                        shutil.move(
                            actual_output_path_from_converter,
                            intended_final_output_path,
                        )
                        logger.info(
                            f"Moved converted file from {actual_output_path_from_converter} to {intended_final_output_path}. RID: {request_id}"
                        )
                    except Exception as e_move:
                        success_conv = False
                        conversion_message_conv = (
                            f"无法将转换后的文件移动到目标位置: {e_move}"
                        )
                        logger.error(
                            f"Failed to move {actual_output_path_from_converter} to {intended_final_output_path}: {e_move}. RID: {request_id}"
                        )
                # else: actual_output_path_from_converter is already the intended_final_output_path, no move needed

                # Re-check success_conv because it might have been set to False during move/check
                if (
                    success_conv
                    and os.path.exists(intended_final_output_path)
                    and os.path.isfile(intended_final_output_path)
                ):
                    processed_files_final.append(
                        {
                            "original_name": original_name,
                            "converted_name": converted_filename,
                            "download_url": reverse(
                                "converter:download_converted_file",
                                args=[
                                    request.user.username,
                                    today_date_str,
                                    converted_filename,
                                ],
                            ),
                            "status": "success",
                            "message": conversion_message_conv or "转换成功。",
                        }
                    )
                    temp_individual_converted_outputs.append(intended_final_output_path)
                else:  # File missing after supposed success
                    success_conv = False
                    conversion_message_conv += " (处理后输出文件丢失)"
                    logger.error(
                        f"pdf_to_file_view: File {intended_final_output_path} missing post-success for {original_name}. RID: {request_id}"
                    )

            if not success_conv:
                processed_files_final.append(
                    {
                        "original_name": original_name,
                        "status": "error",
                        "message": conversion_message_conv or "转换失败",
                    }
                )
                logger.error(
                    f"pdf_to_file_view/{sub_tab}: '{original_name}' failed. Msg: {conversion_message_conv}. OutPath: {actual_output_path_from_converter}. RID: {request_id}"
                )
                if (
                    actual_output_path_from_converter
                    and actual_output_path_from_converter != intended_final_output_path
                    and os.path.exists(actual_output_path_from_converter)
                ):
                    temp_files_to_delete_final.append(actual_output_path_from_converter)
        except Exception as e_conv_pdf_ind:
            logger.error(
                f"pdf_to_file_view: {sub_tab} for {original_name} EXCEPTION: {e_conv_pdf_ind}. RID: {request_id}",
                exc_info=True,
            )
            processed_files_final.append(
                {
                    "original_name": original_name,
                    "status": "error",
                    "message": f"关键错误: {str(e_conv_pdf_ind)}",
                }
            )
            if actual_output_path_from_converter and os.path.exists(
                actual_output_path_from_converter
            ):
                temp_files_to_delete_final.append(actual_output_path_from_converter)
        finally:
            if source_file_path and os.path.exists(source_file_path):
                temp_files_to_delete_final.append(source_file_path)

    if merge_output and len(temp_individual_converted_outputs) > 1:
        first_converted_path = temp_individual_converted_outputs[0]
        merge_output_ext = os.path.splitext(first_converted_path)[1].lower()
        merged_target_filename = f"merged_pdfs_to_{sub_tab.replace('pdfTo','').lower()}_{request_id}{merge_output_ext}"
        merged_target_path = os.path.join(user_converted_dir, merged_target_filename)
        merge_succ = False
        merge_msg = "合并未尝试或失败。"

        successful_original_pdf_paths_for_merge = []
        success_map = (
            {}
        )  # Map original_name to its temp_path for successfully converted files
        for pf in processed_files_final:
            if pf["status"] == "success":
                original_file_detail = next(
                    (
                        uf_info
                        for uf_info in uploaded_files_info
                        if uf_info["name"] == pf["original_name"]
                    ),
                    None,
                )
                if (
                    original_file_detail
                    and os.path.join(user_converted_dir, pf["converted_name"])
                    in temp_individual_converted_outputs
                ):
                    success_map[pf["original_name"]] = original_file_detail["path"]

        for tico_path in temp_individual_converted_outputs:
            found_original_name = next(
                (
                    pf_entry["original_name"]
                    for pf_entry in processed_files_final
                    if pf_entry.get("status") == "success"
                    and os.path.join(
                        user_converted_dir, pf_entry.get("converted_name", "")
                    )
                    == tico_path
                ),
                None,
            )
            if found_original_name and found_original_name in success_map:
                successful_original_pdf_paths_for_merge.append(
                    success_map[found_original_name]
                )
            else:
                logger.warning(
                    f"pdf_to_file_view: Cannot map temp output {tico_path} to original PDF for merge. RID: {request_id}"
                )

        if (
            not successful_original_pdf_paths_for_merge
            or len(successful_original_pdf_paths_for_merge) <= 1
        ):
            logger.warning(
                f"pdf_to_file_view: Merge requested, not enough original PDFs. RID: {request_id}"
            )
        else:
            current_merge_mode = parsed_params.get(
                f"pdf_to_{sub_tab.replace('pdfTo','').lower()}_mode"
            )

            try:
                if sub_tab == "pdfToWord" and merge_output_ext == ".docx":
                    merge_succ, merge_msg = convert_and_merge_pdfs_to_docx(
                        successful_original_pdf_paths_for_merge,
                        merged_target_path,
                        request_id,
                        mode=current_merge_mode,
                    )
                elif sub_tab == "pdfToPpt" and merge_output_ext == ".pptx":
                    merge_succ, merge_msg = convert_and_merge_pdfs_to_pptx(
                        successful_original_pdf_paths_for_merge,
                        merged_target_path,
                        request_id,
                        ppt_creation_mode=current_merge_mode,
                    )
                elif sub_tab == "pdfToTxt" and merge_output_ext == ".txt":
                    merge_succ, merge_msg = convert_and_merge_pdfs_to_txt(
                        successful_original_pdf_paths_for_merge,
                        merged_target_path,
                        request_id,
                        mode=current_merge_mode,
                    )
                else:
                    merge_msg = f"不支持PDF合并为 {merge_output_ext}。"
                    logger.warning(f"pdf_to_file_view: {merge_msg} RID: {request_id}")

                if merge_succ and os.path.exists(merged_target_path):
                    final_list_after_merge = [
                        res for res in processed_files_final if res["status"] == "error"
                    ]
                    final_list_after_merge.append(
                        {
                            "original_name": f'合并的 {sub_tab.replace("pdfTo","")} (来自 {len(successful_original_pdf_paths_for_merge)} 个PDF)',
                            "converted_name": merged_target_filename,
                            "download_url": reverse(
                                "converter:download_converted_file",
                                args=[
                                    request.user.username,
                                    today_date_str,
                                    merged_target_filename,
                                ],
                            ),
                            "status": "success",
                            "message": merge_msg
                            or f"{len(successful_original_pdf_paths_for_merge)} 个PDF成功合并。",
                        }
                    )
                    processed_files_final = final_list_after_merge
                    temp_files_to_delete_final.extend(temp_individual_converted_outputs)
                elif not merge_succ:
                    processed_files_final.append(
                        {
                            "original_name": "合并操作",
                            "status": "error",
                            "message": f"PDF合并为 {sub_tab.replace('pdfTo','')} 失败: {merge_msg}",
                        }
                    )
                    if os.path.exists(merged_target_path):
                        temp_files_to_delete_final.append(merged_target_path)
            except Exception as e_merge_main:
                logger.error(
                    f"pdf_to_file_view: {sub_tab} merge EXCEPTION: {e_merge_main}. RID: {request_id}",
                    exc_info=True,
                )
                processed_files_final.append(
                    {
                        "original_name": "合并操作",
                        "status": "error",
                        "message": f"PDF合并时关键错误: {str(e_merge_main)}",
                    }
                )
                if os.path.exists(merged_target_path):
                    temp_files_to_delete_final.append(merged_target_path)

    elif merge_output and len(temp_individual_converted_outputs) == 1:
        logger.info(
            f"pdf_to_file_view: Merge requested, only one PDF converted. No merge performed. RID: {request_id}"
        )

    cleanup_temp_files(list(set(temp_files_to_delete_final)), request_id)
    end_time_view = time.perf_counter()  # ADDED: End timer
    duration_seconds_view = round(
        end_time_view - start_time_view, 2
    )  # ADDED: Calculate duration

    if (
        any(result["status"] == "error" for result in processed_files_final)
        or not processed_files_final
    ):
        # Consolidate error messages if any, or provide a generic one
        error_messages = [
            res.get("message", "未知错误")
            for res in processed_files_final
            if res["status"] == "error"
        ]
        if not error_messages:
            error_messages.append("处理PDF转换时发生未知错误或没有文件成功处理。")
        logger.error(
            f"pdf_to_file_view: Processing finished with errors. Duration: {duration_seconds_view}s. Errors: {error_messages}. RequestID: {request_id}"
        )
        # Return a single error response if there were issues
        return format_error_response(
            message="; ".join(error_messages),
            request_id=request_id,
            duration_seconds=duration_seconds_view,
        )  # ADDED duration

    final_result_payload = {
        "results": processed_files_final,
        "request_id": request_id,
        "merge_output": merge_output,  # Reflect whether merge was done or intended
        "duration_seconds": duration_seconds_view,  # ADDED duration
    }
    logger.info(
        f"pdf_to_file_view: Processing complete. Duration: {duration_seconds_view}s. Results: {len(processed_files_final)} items. RequestID: {request_id}"
    )
    return format_json_response(
        results=final_result_payload["results"],
        merge_output=final_result_payload["merge_output"],
        request_id=final_result_payload["request_id"],
        duration_seconds=final_result_payload["duration_seconds"],
    )


@login_required
@require_POST
def process_images_view(request):  # This view is now a deprecated placeholder
    request_id = "".join(random.choices(string.ascii_lowercase + string.digits, k=6))
    main_tab = request.POST.get("main_tab", None)
    merge_output = request.POST.get("merge_output", "false").lower() == "true"

    logger.warning(
        f"process_images_view: DEPRECATED endpoint hit. main_tab: {main_tab}, RequestID: {request_id}. Client should use specific API endpoints."
    )

    error_message = "此通用转换接口已弃用。请更新客户端使用新的专用接口。"
    if main_tab == "fileToPdf":
        error_message = "文件转PDF接口已更新至 /api/file-to-pdf/。请更新客户端。"
    elif main_tab == "imgToFile":
        error_message = "图片转文件接口已更新至 /api/img-to-file/。请更新客户端。"
    elif main_tab == "pdfToFile":
        error_message = "PDF转文件接口已更新至 /api/pdf-to-file/。请更新客户端。"

    return format_error_response(
        message=error_message,
        merge_output=merge_output,
        original_item_name=f"Deprecated Call ({main_tab})",
        request_id=request_id,
        http_status=400,
    )


@login_required
def conversion_history_view(request):
    user = request.user
    user_history_base_dir = os.path.join(settings.BASE_DIR, "his_pic", user.username)

    available_dates = []
    if os.path.exists(user_history_base_dir):
        for item in os.listdir(user_history_base_dir):
            if os.path.isdir(os.path.join(user_history_base_dir, item)):
                if len(item) == 8 and item.isdigit():
                    available_dates.append(item)
        available_dates.sort(reverse=True)

    selected_date_str = request.GET.get("date", None)
    converted_files_info = []

    if selected_date_str and selected_date_str in available_dates:
        date_specific_converted_dir = os.path.join(
            user_history_base_dir, selected_date_str, "converted_files"
        )
        if os.path.exists(date_specific_converted_dir):
            for filename in os.listdir(date_specific_converted_dir):
                if filename.endswith(".meta"):  # Skip .meta files themselves
                    continue

                file_path = os.path.join(date_specific_converted_dir, filename)
                if os.path.isfile(file_path):
                    original_name_display = os.path.splitext(filename)[0]  # Fallback
                    meta_file_path = f"{file_path}.meta"
                    if os.path.exists(meta_file_path):
                        try:
                            with open(meta_file_path, "r", encoding="utf-8") as mf:
                                original_name_display = mf.read()
                        except Exception as e:
                            logger.error(
                                f"Error reading .meta file {meta_file_path}: {e}"
                            )

                    download_url = f"{settings.MEDIA_URL}{user.username}/{selected_date_str}/converted_files/{filename}"
                    delete_url = reverse(
                        "converter:delete_converted_file",
                        args=[selected_date_str, filename],
                    )

                    converted_files_info.append(
                        {
                            "original_name": original_name_display,
                            "converted_name": filename,
                            "download_url": download_url,
                            "delete_url": delete_url,  # Use the generated one, not from request
                            "status": "已完成",
                        }
                    )

    context = {
        "available_dates": available_dates,
        "selected_date": selected_date_str,
        "converted_files": converted_files_info,
        "page_title": "历史转换记录",
        "current_nav": "history",
    }
    return render(request, "converter/conversion_history.html", context)


@login_required
@require_POST
def delete_converted_file_view(request, date_str, filename):
    user = request.user
    file_path = os.path.join(
        settings.BASE_DIR,
        "his_pic",
        user.username,
        date_str,
        "converted_files",
        filename,
    )
    meta_file_path = f"{file_path}.meta"

    file_deleted = False
    if os.path.exists(file_path) and os.path.isfile(file_path):
        try:
            os.remove(file_path)
            messages.success(request, f"文件 '{filename}' 已成功删除。")
            logger.info(f"User {user.username} deleted file: {file_path}")
            file_deleted = True

            # Attempt to delete corresponding .meta file
            if os.path.exists(meta_file_path):
                try:
                    os.remove(meta_file_path)
                    logger.info(
                        f"User {user.username} deleted meta file: {meta_file_path}"
                    )
                except OSError as e:
                    logger.warning(
                        f"Error deleting meta file {meta_file_path} for user {user.username}: {e}"
                    )

            # Check if the converted_files directory is now empty
            converted_dir_path = os.path.dirname(file_path)
            if not os.listdir(converted_dir_path):
                try:
                    os.rmdir(converted_dir_path)
                    logger.info(f"Removed empty directory: {converted_dir_path}")
                    # Check if the parent date directory is now empty (uploads might still be there)
                    date_dir_path = os.path.dirname(converted_dir_path)
                    # We only remove the date dir if both 'uploads' and 'converted_files' are gone or empty
                    uploads_dir_path = os.path.join(date_dir_path, "uploads")
                    can_delete_date_dir = True
                    if os.path.exists(uploads_dir_path) and os.listdir(
                        uploads_dir_path
                    ):
                        can_delete_date_dir = False

                    if not os.path.exists(converted_dir_path) and not os.path.exists(
                        uploads_dir_path
                    ):  # both gone
                        pass  # can delete
                    elif (
                        not os.path.exists(converted_dir_path)
                        and os.path.exists(uploads_dir_path)
                        and not os.listdir(uploads_dir_path)
                    ):  # converted gone, uploads empty
                        os.rmdir(uploads_dir_path)  # remove empty uploads
                        logger.info(f"Removed empty directory: {uploads_dir_path}")
                    elif (
                        can_delete_date_dir
                    ):  # converted was removed, uploads never existed or was already removed
                        pass
                    else:  # uploads still has content or converted_files was not empty
                        can_delete_date_dir = False

                    if can_delete_date_dir and not os.listdir(
                        date_dir_path
                    ):  # Check if date_dir is truly empty now
                        os.rmdir(date_dir_path)
                        logger.info(f"Removed empty date directory: {date_dir_path}")

                except OSError as e:
                    logger.error(
                        f"Error removing directory for user {user.username} after file deletion: {e}"
                    )
                    # Don't send this specific error to user, file deletion was successful.

        except OSError as e:
            messages.error(request, f"删除文件 '{filename}' 时出错: {e}")
            logger.error(
                f"Error deleting file {file_path} for user {user.username}: {e}"
            )
    else:
        messages.error(request, "文件未找到或无法删除。")
        logger.warning(
            f"Attempt to delete non-existent file by {user.username}: {file_path}"
        )

    # Redirect to the history page, potentially without the date if the folder was removed
    # Or always redirect to the general history page to show the date is gone from the list
    return redirect(reverse("converter:conversion_history"))


@login_required
@require_POST  # Ensure this view is only accessed via POST
def delete_all_for_date_view(request, date_str):
    user = request.user
    user_date_dir = os.path.join(settings.BASE_DIR, "his_pic", user.username, date_str)

    if not os.path.exists(user_date_dir) or not os.path.isdir(user_date_dir):
        messages.error(request, f"日期 '{date_str}' 的记录不存在或无法访问。")
        return redirect(reverse("converter:conversion_history") + f"?date={date_str}")

    converted_files_dir = os.path.join(user_date_dir, "converted_files")
    uploads_dir = os.path.join(user_date_dir, "uploads")

    deleted_something = False
    try:
        # Delete files in converted_files directory
        if os.path.exists(converted_files_dir):
            for item_name in os.listdir(converted_files_dir):
                item_path = os.path.join(converted_files_dir, item_name)
                try:
                    if os.path.isfile(item_path) or os.path.islink(item_path):
                        os.remove(item_path)
                        logger.info(
                            f"User {user.username} deleted file/link during mass delete: {item_path}"
                        )
                    elif os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                        logger.info(
                            f"User {user.username} deleted directory during mass delete: {item_path}"
                        )
                    deleted_something = True
                except OSError as e:
                    logger.warning(
                        f"Error deleting item {item_path} during mass delete for user {user.username}: {e}"
                    )
                    messages.warning(
                        request, f"删除 '{item_name}' 时出错，但会继续尝试。"
                    )
            # Attempt to remove the converted_files directory if empty
            if not os.listdir(
                converted_files_dir
            ):  # Should be empty if all items were deleted
                os.rmdir(converted_files_dir)
                logger.info(f"Removed empty directory: {converted_files_dir}")

        # Delete files in uploads directory
        if os.path.exists(uploads_dir):
            for item_name in os.listdir(uploads_dir):
                item_path = os.path.join(uploads_dir, item_name)
                try:
                    if os.path.isfile(item_path) or os.path.islink(item_path):
                        os.remove(item_path)
                        logger.info(
                            f"User {user.username} deleted uploaded file/link during mass delete: {item_path}"
                        )
                    elif os.path.isdir(item_path):
                        shutil.rmtree(item_path)
                        logger.info(
                            f"User {user.username} deleted uploaded directory during mass delete: {item_path}"
                        )
                    deleted_something = True
                except OSError as e:
                    logger.warning(
                        f"Error deleting uploaded item {item_path} during mass delete for user {user.username}: {e}"
                    )
                    messages.warning(
                        request,
                        f"删除上传的文件/目录 '{item_name}' 时出错，但会继续尝试。",
                    )
            # Attempt to remove the uploads directory if empty
            if not os.listdir(uploads_dir):
                os.rmdir(uploads_dir)
                logger.info(f"Removed empty directory: {uploads_dir}")

        # Attempt to remove the date directory itself if it's now empty
        if not os.listdir(user_date_dir):
            os.rmdir(user_date_dir)
            logger.info(f"Removed empty date directory: {user_date_dir}")
            messages.success(request, f"日期 '{date_str}' 的所有记录已成功清除。")
        elif deleted_something:
            messages.success(
                request,
                f"日期 '{date_str}' 的部分或全部文件已清除。可能仍有空目录结构残留。",
            )
        else:
            messages.info(request, f"日期 '{date_str}' 下没有找到可清除的文件。")

    except OSError as e:
        messages.error(request, f"清除日期 '{date_str}' 的记录时发生错误: {e}")
        logger.error(
            f"Error during mass delete for user {user.username}, date {date_str}: {e}",
            exc_info=True,
        )

    # Redirect to the history page, potentially without the date if the folder was removed
    # Or always redirect to the general history page to show the date is gone from the list
    return redirect(reverse("converter:conversion_history"))


@login_required
def download_converted_file_view(request, username, date_str, filename):
    # Security check: Ensure the logged-in user matches the username in the URL
    # or the logged-in user is a superuser.
    if not (request.user.username == username or request.user.is_superuser):
        logger.warning(
            f"Permission denied for user {request.user.username} trying to download file for user {username}."
        )
        raise PermissionDenied("您没有权限下载此文件。")

    # Construct the full path to the file
    # Ensure to use settings.BASE_DIR or another secure base path for `his_pic`
    file_path = os.path.join(
        settings.BASE_DIR, "his_pic", username, date_str, "converted_files", filename
    )

    logger.debug(
        f"Attempting to serve file: {file_path} for user {request.user.username}"
    )

    if os.path.exists(file_path):
        try:
            return FileResponse(
                open(file_path, "rb"), as_attachment=True, filename=filename
            )
        except Exception as e:
            logger.error(f"Error serving file {file_path}: {e}", exc_info=True)
            raise Http404("下载文件时发生错误。")
    else:
        logger.error(
            f"File not found for download by {request.user.username}: {file_path}"
        )
        raise Http404("文件未找到。")


@csrf_exempt  # Ensure CSRF exemption if you test directly without a form including {% csrf_token %}
@require_POST
def process_video_extraction_view(request):
    request_id = "".join(
        random.choices(string.ascii_lowercase + string.digits, k=10)
    )  # Unique ID for this request
    today_date_str = datetime.now().strftime("%Y%m%d")
    logger.info(
        f"process_video_extraction_view: Received request. RequestID: {request_id}"
    )

    user_upload_dir, user_converted_dir = "", ""
    try:
        user_upload_dir, user_converted_dir = ensure_user_directories(
            request.user.username, today_date_str
        )
    except Exception as e:
        logger.critical(
            f"process_video_extraction_view: Failed to create user directories for {request.user.username}. Error: {e}. RequestID: {request_id}",
            exc_info=True,
        )
        return format_error_response(
            message="服务器错误：无法创建用户目录。",
            merge_output=False,  # 视频处理不涉及合并输出
            request_id=request_id,
        )

    video_file_obj = request.FILES.get(
        "video_file"
    )  # 与前端 formData.append('video_file', uploadedVideoFile) 对应
    scene_threshold_str = request.POST.get(
        "scene_detection_threshold", "10.0"
    )  # 与前端 formData.append('scene_detection_threshold', ...) 对应
    group_size_str = request.POST.get(
        "deduplication_group_size", "5"
    )  # 与前端 formData.append('deduplication_group_size', ...) 对应

    if not video_file_obj:
        logger.warning(
            f"process_video_extraction_view: No video file uploaded. RequestID: {request_id}"
        )
        return format_error_response(
            message="没有上传视频文件。",
            merge_output=False,  # 视频处理不涉及合并输出
            request_id=request_id,
        )
    try:
        scene_threshold = float(scene_threshold_str)
        group_size = int(group_size_str)
    except ValueError:
        logger.warning(
            f"process_video_extraction_view: Invalid threshold or group size. T: {scene_threshold_str}, G: {group_size_str}. RequestID: {request_id}"
        )
        return format_error_response(
            message="场景阈值或分组大小参数无效。",
            merge_output=False,  # 视频处理不涉及合并输出
            request_id=request_id,
        )

    temp_video_path, original_video_filename, safe_video_filename = save_uploaded_file(
        video_file_obj, user_upload_dir, request_id
    )
    if not temp_video_path:
        logger.error(
            f"process_video_extraction_view: Failed to save uploaded video file: {original_video_filename}. RequestID: {request_id}"
        )
        return format_error_response(
            message=f'视频文件 "{original_video_filename}" 上传保存失败。',
            merge_output=False,  # 视频处理不涉及合并输出
            request_id=request_id,
        )

    # Determine script path relative to settings.BASE_DIR (extract_web)
    # settings.BASE_DIR is .../extract_doc/extract_web
    # script is in .../extract_doc/
    script_base_dir = os.path.abspath(os.path.join(settings.BASE_DIR, ".."))
    script_path = os.path.join(script_base_dir, "extract_video_snapshots.py")

    if not os.path.exists(script_path):
        logger.error(
            f"process_video_extraction_view: Snapshot script not found at {script_path}. RequestID: {request_id}"
        )
        # Use the 'type' key for cleanup_temp_files if passing a list of dicts
        cleanup_temp_files([{"path": temp_video_path, "type": "file"}], request_id)
        return format_error_response(
            message="服务器配置错误：找不到视频处理脚本。",
            merge_output=False,  # 视频处理不涉及合并输出
            request_id=request_id,
        )

    # Create a temporary directory for script execution
    exec_temp_dir = os.path.join(user_upload_dir, f"video_exec_{request_id}")
    os.makedirs(exec_temp_dir, exist_ok=True)

    # Define final target directories for snapshots in user's history
    # These names now include safe_video_filename and request_id for better uniqueness
    target_raw_snapshots_dir = os.path.join(
        user_converted_dir, f"video-snapshot_raw_{safe_video_filename}_{request_id}"
    )
    target_dedup_snapshots_dir = os.path.join(
        user_converted_dir, f"video-snapshot_dedup_{safe_video_filename}_{request_id}"
    )
    # These directories will be created by shutil.copytree later if source exists.

    def stream_video_processing_response(
        _script_path_arg,
        _temp_video_path_arg,
        _exec_temp_dir_arg,
        _scene_threshold_arg,
        _group_size_arg,
        _original_video_filename_arg,
        _safe_video_filename_arg,
        _user_converted_dir_arg,
        _today_date_str_arg,
        _target_raw_snapshots_dir_arg,
        _target_dedup_snapshots_dir_arg,
        _request_user_username_arg,
        _request_id_arg,
    ):
        start_time_stream = (
            time.perf_counter()
        )  # ADDED: Start timer for the stream processing
        # Files and dirs to clean up at the end of this specific stream
        _temp_files_to_clean_stream_arg = [
            {"path": _temp_video_path_arg, "type": "file"},
            {"path": _exec_temp_dir_arg, "type": "dir"},
        ]
        process_stream = None
        final_result_payload_stream_var = {}

        try:
            command_list_stream = [
                sys.executable,
                _script_path_arg,
                "--video_file",
                _temp_video_path_arg,  # MODIFIED: Changed from --video_file_path
                "--output_base_dir",
                _exec_temp_dir_arg,
                "--threshold",
                str(
                    _scene_threshold_arg
                ),  # MODIFIED: Changed from --scene_detection_thresh
                "--group_size",
                str(
                    _group_size_arg
                ),  # MODIFIED: Changed from --deduplication_group_size
            ]
            logger.info(
                f"stream_video_processing_response: Executing video script for {_original_video_filename_arg} (RequestID: {_request_id_arg}): {' '.join(command_list_stream)}"
            )

            yield f"data: {json.dumps({'type': 'info', 'message': '视频处理脚本已启动，请稍候...'})}\n\n"

            process_stream = subprocess.Popen(
                command_list_stream,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                encoding="utf-8",
                errors="replace",
                bufsize=1,
                universal_newlines=True,
                env=os.environ.copy(),  # ADDED: Pass a copy of the current environment
            )

            if process_stream and process_stream.poll() is None:
                if process_stream.stderr:
                    for line_from_stderr_stream in iter(
                        process_stream.stderr.readline, ""
                    ):
                        original_line_stream = line_from_stderr_stream.strip()
                        if not original_line_stream:
                            continue

                        # Prepare for logging to console (e.g., GBK) - Stricter ASCII enforcement for logging
                        ascii_safe_log_line = original_line_stream.encode(
                            "ascii", "ignore"
                        ).decode("ascii")
                        logger.debug(
                            f"Script STDERR line (ASCII for console): {ascii_safe_log_line}. RequestID: {_request_id_arg}"
                        )

                        cleaned_display_line_stream = (
                            original_line_stream.encode("ascii", "ignore")
                            .decode("utf-8", "ignore")
                            .replace("\\\\ufffd", "")
                            .strip()
                        )
                        progress_match_stream = re.search(
                            r"(\d+)/(\d+)\\s*\\((.*?)%\\)", original_line_stream
                        )  # Regex on original_line_stream for robustness
                        if progress_match_stream:
                            _, _, percent_str_stream = progress_match_stream.groups()
                            try:
                                percent_stream = float(percent_str_stream)
                                yield f"data: {json.dumps({'type': 'progress', 'percent': percent_stream, 'text': cleaned_display_line_stream})}\n\n"
                            except ValueError:
                                yield f"data: {json.dumps({'type': 'info', 'message': cleaned_display_line_stream})}\n\n"
                        else:
                            yield f"data: {json.dumps({'type': 'info', 'message': cleaned_display_line_stream})}\n\n"

            stdout_data_collected_stream = ""
            stderr_data_remaining_collected_stream = ""
            return_code_collected_stream = None

            try:
                if process_stream:
                    if process_stream.poll() is None:  # Check if it hasn't finished
                        logger.info(
                            f"Waiting for script process to complete using communicate(). PID: {process_stream.pid}. RequestID: {_request_id_arg}"
                        )
                        # Timeout for communicate should be for script's finalization (stdout, exit),
                        # as stderr progress loop is done. 10 minutes for finalization.
                        stdout_data, stderr_data = process_stream.communicate(
                            timeout=600
                        )
                        stdout_data_collected_stream = stdout_data
                        # The loop consumed most/all stderr for progress. This captures any remaining.
                        stderr_data_remaining_collected_stream = stderr_data

                    return_code_collected_stream = process_stream.returncode
                    logger.info(
                        f"Script process (after communicate) ended with return code: {return_code_collected_stream}. PID: {process_stream.pid if process_stream else 'N/A'}. RequestID: {_request_id_arg}"
                    )
                else:
                    logger.error(
                        f"Process was not started (is None) for {_request_id_arg} in stream. Cannot wait or get outputs."
                    )
                    return_code_collected_stream = -10  # Special code for Popen failure
            except subprocess.TimeoutExpired:
                logger.error(
                    f"Script process communicate() TIMED OUT after 600s. Killing process. PID: {process_stream.pid if process_stream else 'N/A'}. RequestID: {_request_id_arg}"
                )
                if process_stream:
                    process_stream.kill()
                    try:
                        # Attempt to get any final output after kill
                        stdout_after_kill_stream, stderr_after_kill_stream = (
                            process_stream.communicate(timeout=10)
                        )
                        # Append to existing rather than replace, in case some data was already present before timeout logic
                        if stdout_after_kill_stream:
                            stdout_data_collected_stream += stdout_after_kill_stream
                        if stderr_after_kill_stream:
                            stderr_data_remaining_collected_stream += (
                                stderr_after_kill_stream
                            )
                        logger.info(
                            f"Script output after kill (PID: {process_stream.pid}): STDOUT='{stdout_data_collected_stream[:200]}...', STDERR='{stderr_data_remaining_collected_stream[:200]}...'. RequestID: {_request_id_arg}"
                        )
                    except subprocess.TimeoutExpired:
                        logger.warning(
                            f"Communicate timed out after killing process (PID: {process_stream.pid}). RequestID: {_request_id_arg}"
                        )
                    except Exception as e_comm_kill_stream:
                        logger.error(
                            f"Exception during communicate after kill (PID: {process_stream.pid}): {e_comm_kill_stream}. RequestID: {_request_id_arg}"
                        )
                return_code_collected_stream = -9  # Special return code for timeout
                final_result_payload_stream_var = {
                    "type": "error",
                    "message": "视频处理脚本执行最终化步骤超时（超过10分钟限制）。",
                    "request_id": _request_id_arg,
                }
                yield f"data: {json.dumps(final_result_payload_stream_var)}\n\n"

            if stdout_data_collected_stream:
                console_safe_stdout = stdout_data_collected_stream.encode(
                    sys.stdout.encoding or "gbk", "replace"
                ).decode(sys.stdout.encoding or "gbk", "ignore")
                logger.info(
                    f"Script STDOUT (final accumulated for {_request_id_arg}):\\n{console_safe_stdout}"
                )
            if stderr_data_remaining_collected_stream:
                # Stricter ASCII enforcement for final stderr log as well
                ascii_safe_remaining_stderr = (
                    stderr_data_remaining_collected_stream.encode(
                        "ascii", "ignore"
                    ).decode("ascii")
                )
                logger.error(
                    f"Script STDERR (final accumulated, ASCII for console for {_request_id_arg}):\\n{ascii_safe_remaining_stderr}"
                )

            if (
                return_code_collected_stream != -9
                and return_code_collected_stream != -10
            ):  # Not a timeout and not a Popen failure
                if return_code_collected_stream == 0:  # Script success
                    logger.info(
                        f"Script executed successfully for {_original_video_filename_arg}. RequestID: {_request_id_arg}"
                    )
                    raw_count_from_stdout = 0
                    dedup_count_from_stdout = 0
                    if stdout_data_collected_stream:
                        raw_match_stdout = re.search(
                            r"Raw snapshots count: (\d+)", stdout_data_collected_stream
                        )
                        if raw_match_stdout:
                            try:
                                raw_count_from_stdout = int(raw_match_stdout.group(1))
                            except ValueError:
                                logger.warning(
                                    f"Could not parse raw_count from stdout. RID: {_request_id_arg}"
                                )
                        dedup_match_stdout = re.search(
                            r"Deduplicated snapshots count: (\d+)",
                            stdout_data_collected_stream,
                        )
                        if dedup_match_stdout:
                            try:
                                dedup_count_from_stdout = int(
                                    dedup_match_stdout.group(1)
                                )
                            except ValueError:
                                logger.warning(
                                    f"Could not parse dedup_count from stdout. RID: {_request_id_arg}"
                                )

                    source_raw_dir_script_output = os.path.join(
                        _exec_temp_dir_arg, "video-snapshot"
                    )
                    source_dedup_dir_script_output = os.path.join(
                        _exec_temp_dir_arg, "video-snapshot-duplicate"
                    )

                    current_results_list_for_payload = []

                    if os.path.exists(source_raw_dir_script_output) and os.path.isdir(
                        source_raw_dir_script_output
                    ):
                        if os.path.exists(_target_raw_snapshots_dir_arg):
                            shutil.rmtree(_target_raw_snapshots_dir_arg)
                        shutil.copytree(
                            source_raw_dir_script_output,
                            _target_raw_snapshots_dir_arg,
                            dirs_exist_ok=False,
                        )
                        logger.info(
                            f"Copied raw snapshots to {_target_raw_snapshots_dir_arg}. RequestID: {_request_id_arg}"
                        )

                        raw_zip_base_name_for_payload = os.path.join(
                            _user_converted_dir_arg,
                            f"raw_frames_{_safe_video_filename_arg}_{_request_id_arg}",
                        )
                        raw_zip_file_path_for_payload = shutil.make_archive(
                            raw_zip_base_name_for_payload,
                            "zip",
                            _target_raw_snapshots_dir_arg,
                        )
                        raw_zip_filename_for_payload = os.path.basename(
                            raw_zip_file_path_for_payload
                        )
                        raw_meta_file_path_for_payload = (
                            f"{raw_zip_file_path_for_payload}.meta"
                        )
                        try:
                            with open(
                                raw_meta_file_path_for_payload, "w", encoding="utf-8"
                            ) as mf_raw_stream:
                                mf_raw_stream.write(_original_video_filename_arg)
                            logger.info(
                                f"Created .meta for raw ZIP: {raw_meta_file_path_for_payload}. RequestID: {_request_id_arg}"
                            )
                        except Exception as e_meta_raw_stream:
                            logger.error(
                                f"Failed .meta for raw ZIP {raw_meta_file_path_for_payload}: {e_meta_raw_stream}. RequestID: {_request_id_arg}"
                            )
                        current_results_list_for_payload.append(
                            {
                                "original_name": f"{_original_video_filename_arg} (原始截图)",
                                "converted_name": raw_zip_filename_for_payload,
                                "download_url": reverse(
                                    "converter:download_converted_file",
                                    args=[
                                        _request_user_username_arg,
                                        _today_date_str_arg,
                                        raw_zip_filename_for_payload,
                                    ],
                                ),
                                "status": "success",
                                "message": f"包含所有原始提取的截图 ({raw_count_from_stdout} 张)。",
                            }
                        )
                    else:
                        logger.warning(
                            f"Raw snapshot output directory not found: {source_raw_dir_script_output}. RequestID: {_request_id_arg}"
                        )

                    if os.path.exists(source_dedup_dir_script_output) and os.path.isdir(
                        source_dedup_dir_script_output
                    ):
                        if os.path.exists(_target_dedup_snapshots_dir_arg):
                            shutil.rmtree(_target_dedup_snapshots_dir_arg)
                        shutil.copytree(
                            source_dedup_dir_script_output,
                            _target_dedup_snapshots_dir_arg,
                            dirs_exist_ok=False,
                        )
                        logger.info(
                            f"Copied deduplicated snapshots to {_target_dedup_snapshots_dir_arg}. RequestID: {_request_id_arg}"
                        )

                        dedup_zip_base_name_for_payload = os.path.join(
                            _user_converted_dir_arg,
                            f"deduplicated_frames_{_safe_video_filename_arg}_{_request_id_arg}",
                        )
                        dedup_zip_file_path_for_payload = shutil.make_archive(
                            dedup_zip_base_name_for_payload,
                            "zip",
                            _target_dedup_snapshots_dir_arg,
                        )
                        dedup_zip_filename_for_payload = os.path.basename(
                            dedup_zip_file_path_for_payload
                        )
                        dedup_meta_file_path_for_payload = (
                            f"{dedup_zip_file_path_for_payload}.meta"
                        )
                        try:
                            with open(
                                dedup_meta_file_path_for_payload, "w", encoding="utf-8"
                            ) as mf_dedup_stream:
                                mf_dedup_stream.write(_original_video_filename_arg)
                            logger.info(
                                f"Created .meta for deduplicated ZIP: {dedup_meta_file_path_for_payload}. RequestID: {_request_id_arg}"
                            )
                        except Exception as e_meta_dedup_stream:
                            logger.error(
                                f"Failed .meta for dedup ZIP {dedup_meta_file_path_for_payload}: {e_meta_dedup_stream}. RequestID: {_request_id_arg}"
                            )
                        current_results_list_for_payload.append(
                            {
                                "original_name": f"{_original_video_filename_arg} (去重截图)",
                                "converted_name": dedup_zip_filename_for_payload,
                                "download_url": reverse(
                                    "converter:download_converted_file",
                                    args=[
                                        _request_user_username_arg,
                                        _today_date_str_arg,
                                        dedup_zip_filename_for_payload,
                                    ],
                                ),
                                "status": "success",
                                "message": f"视频帧去重完成。原始截图: {raw_count_from_stdout} 张，去重后截图: {dedup_count_from_stdout} 张。",
                            }
                        )
                    else:
                        logger.warning(
                            f"Deduplicated snapshot output directory not found: {source_dedup_dir_script_output}. RequestID: {_request_id_arg}"
                        )
                        if not any(
                            r["status"] == "success"
                            and "原始截图" in r["original_name"]
                            for r in current_results_list_for_payload
                        ):
                            current_results_list_for_payload.append(
                                {
                                    "original_name": _original_video_filename_arg,
                                    "status": "error",
                                    "message": "脚本成功但未找到去重截图。",
                                }
                            )

                    if (
                        not current_results_list_for_payload
                    ):  # Both raw and dedup dirs were not found
                        logger.error(
                            f"Neither raw nor dedup output found. RawSrc: {source_raw_dir_script_output}, DedupSrc: {source_dedup_dir_script_output}. RID: {_request_id_arg}"
                        )
                        current_results_list_for_payload.append(
                            {
                                "original_name": _original_video_filename_arg,
                                "status": "error",
                                "message": "脚本成功但未找到任何截图输出目录。",
                            }
                        )
                    final_result_payload_stream_var = {
                        "type": "result",
                        "results": current_results_list_for_payload,
                        "request_id": _request_id_arg,
                        "merge_output": False,
                    }

                elif (
                    return_code_collected_stream is not None
                ):  # Script failed (not 0, not -9, not -10)
                    logger.error(
                        f"stream_video_processing_response: Script failed with code {return_code_collected_stream}. Input: {_original_video_filename_arg}. RID: {_request_id_arg}"
                    )
                    error_detail = (
                        stderr_data_remaining_collected_stream[:500]
                        if stderr_data_remaining_collected_stream
                        else "(无详细错误信息)"
                    )
                    if (
                        not stderr_data_remaining_collected_stream
                        and stdout_data_collected_stream
                    ):
                        error_detail += f" STDOUT: {stdout_data_collected_stream[:300]}"
                    final_result_payload_stream_var = {
                        "type": "error",
                        "message": f"视频处理脚本执行失败: {error_detail}",
                        "request_id": _request_id_arg,
                    }

            elif return_code_collected_stream == -10:  # Popen failure
                logger.error(
                    f"stream_video_processing_response: Popen failed. Input: {_original_video_filename_arg}. RID: {_request_id_arg}"
                )
                final_result_payload_stream_var = {
                    "type": "error",
                    "message": "无法启动视频处理脚本。",
                    "request_id": _request_id_arg,
                }
            # Ensure final_result_payload_stream_var is initialized if no specific path above set it (e.g., if it was a timeout handled earlier)
            if (
                not final_result_payload_stream_var
                and return_code_collected_stream == -9
            ):  # Was a timeout and already yielded error
                # In case of timeout, final_result_payload_stream_var might have been set in the timeout exception block.
                # If not, we might just send a final marker or rely on the previously yielded timeout message.
                # For now, let's assume the timeout block correctly set it or a final yield is not strictly needed beyond that.
                pass  # Or set a default if necessary, but timeout already yielded.
            elif (
                not final_result_payload_stream_var
            ):  # Fallback for any other unexpected unhandled case
                logger.error(
                    f"stream_video_processing_response: final_result_payload_stream_var not set. Input: {_original_video_filename_arg}. RID: {_request_id_arg}. ReturnCode: {return_code_collected_stream}"
                )
                final_result_payload_stream_var = {
                    "type": "error",
                    "message": "视频处理时发生未知服务端内部错误。",
                    "request_id": _request_id_arg,
                }

        except Exception as e_stream_outer:
            logger.error(
                f"stream_video_processing_response: Outer exception: {e_stream_outer}. Traceback: {traceback.format_exc()}. RID: {_request_id_arg}"
            )
            final_result_payload_stream_var = {
                "type": "error",
                "message": f"视频流处理过程中发生意外错误: {e_stream_outer}",
                "request_id": _request_id_arg,
            }

        finally:
            end_time_stream = (
                time.perf_counter()
            )  # ADDED: End timer for the stream processing
            duration_seconds_stream = round(
                end_time_stream - start_time_stream, 2
            )  # ADDED: Calculate duration
            logger.info(
                f"stream_video_processing_response: Finalizing stream. Duration: {duration_seconds_stream}s. RequestID: {_request_id_arg}"
            )

            if final_result_payload_stream_var:  # Ensure it's not empty
                final_result_payload_stream_var["duration_seconds"] = (
                    duration_seconds_stream  # ADDED duration to final payload
                )
                yield f"data: {json.dumps(final_result_payload_stream_var)}\n\n"
            else:
                # This case should ideally be avoided by ensuring final_result_payload_stream_var is always set.
                # However, as a fallback, yield a generic error with duration.
                fallback_error_payload = {
                    "type": "error",
                    "message": "视频处理结束，但未生成明确结果。请检查日志。",
                    "request_id": _request_id_arg,
                    "duration_seconds": duration_seconds_stream,
                }
                logger.warning(
                    f"stream_video_processing_response: final_result_payload_stream_var was empty at the end. Yielding fallback. RID: {_request_id_arg}"
                )
                yield f"data: {json.dumps(fallback_error_payload)}\n\n"

            # This is the critical part for cleanup after streaming response is fully sent.
            # Schedule cleanup of temporary files associated with this specific stream processing.
            # cleanup_temp_files(_temp_files_to_clean_stream_arg, _request_id_arg)
            # logger.info(f"stream_video_processing_response: Scheduled cleanup for stream-specific temp files. RID: {_request_id_arg}")
            # Commented out cleanup as it's handled in the main view after generator is exhausted.

    # ... (rest of process_video_extraction_view)
    # The main view will call the generator and then perform global cleanup.
    # Create the StreamingHttpResponse with the generator.
    response = StreamingHttpResponse(
        stream_video_processing_response(
            script_path,
            temp_video_path,
            exec_temp_dir,
            scene_threshold,
            group_size,
            original_video_filename,
            safe_video_filename,
            user_converted_dir,
            today_date_str,
            target_raw_snapshots_dir,
            target_dedup_snapshots_dir,
            request.user.username,
            request_id,
        ),
        content_type="text/event-stream",
    )

    # Perform cleanup AFTER the streaming response has finished (or generator is exhausted)
    # This is tricky as the response is returned before the generator is fully consumed.
    # A better way for post-stream cleanup might involve signals or a different architecture.
    # For now, we rely on the main view's structure or a later scheduled task if needed.
    # The _temp_files_to_clean_stream_arg in the generator is a good start if we can pass its state out.
    # The exec_temp_dir is created by the script and should be cleaned up by it, or explicitly here.
    # For now, this exec_temp_dir (created as os.path.join(user_upload_dir, f"video_exec_{request_id}"))
    # will be cleaned up IF the stream_video_processing_response is fully iterated by the client and server.
    # If connection drops, it might not be. This is a general challenge with external processes and temp dirs.

    # Add items to the main view's cleanup list
    # temp_files_to_clean_main_view.append({'path': temp_video_path, 'type': 'file'}) # Already added from save_uploaded_file
    # temp_files_to_clean_main_view.append({'path': exec_temp_dir, 'type': 'dir'}) # This is specific to video exec
    # This exec_temp_dir needs to be cleaned. The stream_video_processing_response's finally block is too early if using StreamingHttpResponse.
    # A more robust solution would be to ensure this is cleaned perhaps by the script itself or a managing task.
    # For now, this exec_temp_dir (created as os.path.join(user_upload_dir, f"video_exec_{request_id}"))
    # will be cleaned up IF the stream_video_processing_response is fully iterated by the client and server.
    # If connection drops, it might not be. This is a general challenge with external processes and temp dirs.

    return response


# Celery task status check view (if you integrate Celery later)
@login_required
def check_task_status_view(request, task_id):
    # Placeholder implementation
    # In a real scenario, you would query your task queue (e.g., Celery) for the task status.
    logger.info(
        f"Checking status for task_id: {task_id}. User: {request.user.username}"
    )

    # Simulate some possible states
    # This is highly dependent on how you implement tasks
    if task_id.startswith("sim_success_"):
        return JsonResponse(
            {
                "task_id": task_id,
                "status": "SUCCESS",
                "result": {
                    "message": "Task completed successfully!",
                    "output_url": "/media/dummy_output.zip",
                },
            }
        )
    elif task_id.startswith("sim_pending_"):
        return JsonResponse(
            {
                "task_id": task_id,
                "status": "PENDING",
                "result": {"message": "Task is waiting to be processed."},
            }
        )
    elif task_id.startswith("sim_processing_"):
        return JsonResponse(
            {
                "task_id": task_id,
                "status": "PROCESSING",
                "result": {
                    "message": "Task is currently being processed.",
                    "progress": 50,
                },
            }
        )
    elif task_id.startswith("sim_failure_"):
        return JsonResponse(
            {
                "task_id": task_id,
                "status": "FAILURE",
                "result": {"message": "Task failed to complete."},
            }
        )
    else:
        # Default: Simulate task not found or still processing for a generic ID
        # You might want to return a 404 if the task ID is definitively not found
        logger.warning(
            f"Task ID {task_id} not found or status unknown (placeholder). Returning as PENDING."
        )
        return JsonResponse(
            {
                "task_id": task_id,
                "status": "PENDING",
                "message": "Status unknown or task not found (placeholder response).",
            }
        )


@login_required
@require_POST
def speech_to_text_view(request):
    request_id = (
        generate_request_id()
    )  # Assuming generate_request_id is globally available or import it
    logger.info(f"speech_to_text_view: Received request. RequestID: {request_id}")
    logger.info(
        f"speech_to_text_view: Request Headers: {request.headers}. RequestID: {request_id}"
    )
    logger.info(
        f"speech_to_text_view: Raw Request Body: {request.body[:500]}. RequestID: {request_id}"
    )

    start_time_view = time.perf_counter()

    try:
        if not request.body:
            logger.warning(
                f"speech_to_text_view: Empty request body. RequestID: {request_id}"
            )
            return format_error_response(
                message="Request body is empty.",
                merge_output=False,
                http_status=400,
                request_id=request_id,
            )

        data = json.loads(request.body.decode("utf-8"))
        audio_url = data.get("audio_url")
        # MODIFIED: Expect 'hotwords_config' instead of 'hotwords'
        hotwords_config_from_request = data.get("hotwords_config")

        if not audio_url:
            logger.warning(
                f"speech_to_text_view: 'audio_url' not found in request. RequestID: {request_id}, Data: {data}"
            )
            return format_error_response(
                message="'audio_url' is required in the request body.",
                merge_output=False,
                http_status=400,
                request_id=request_id,
            )

        # Validate hotwords_config_from_request if it exists
        if hotwords_config_from_request is not None:
            if not isinstance(hotwords_config_from_request, list):
                logger.warning(
                    f"speech_to_text_view: 'hotwords_config' parameter must be a list. RequestID: {request_id}, Received: {hotwords_config_from_request}"
                )
                return format_error_response(
                    message="'hotwords_config' must be a list of hotword definitions.",
                    merge_output=False,
                    http_status=400,
                    request_id=request_id,
                )
            for item in hotwords_config_from_request:
                if not (
                    isinstance(item, dict)
                    and "text" in item
                    and isinstance(item["text"], str)
                    and "weight" in item
                    and isinstance(item["weight"], int)
                    and "lang" in item
                    and isinstance(item["lang"], str)
                ):
                    logger.warning(
                        f"speech_to_text_view: Invalid item in 'hotwords_config': {item}. Each item must be a dict with str 'text', int 'weight', and str 'lang'. RequestID: {request_id}"
                    )
                    return format_error_response(
                        message="Invalid structure for item in 'hotwords_config'.",
                        merge_output=False,
                        http_status=400,
                        request_id=request_id,
                    )
            logger.info(
                f"speech_to_text_view: Received hotwords_config: {hotwords_config_from_request}. RequestID: {request_id}"
            )
        else:
            logger.info(
                f"speech_to_text_view: No hotwords_config provided. RequestID: {request_id}"
            )

        logger.info(
            f"speech_to_text_view: Processing URL: {audio_url}. RequestID: {request_id}"
        )

        # 文件大小校验（最大500MB）
        try:
            import requests

            head_resp = requests.head(audio_url, timeout=10)
            content_length = int(head_resp.headers.get("Content-Length", 0))
            max_size = 500 * 1024 * 1024  # 500MB
            if content_length > max_size:
                logger.warning(
                    f"speech_to_text_view: 文件大小超出限制 ({content_length} bytes > 500MB). RequestID: {request_id}"
                )
                return format_error_response(
                    message="音频文件大小不能超过500MB。",
                    merge_output=False,
                    http_status=400,
                    request_id=request_id,
                )
        except Exception as e:
            logger.warning(
                f"speech_to_text_view: 获取音频文件大小失败，无法校验。RequestID: {request_id}, Error: {e}"
            )

        # Pass hotwords_config_from_request to the transcription function
        transcription_result = transcribe_audio_dashscope(
            audio_url, hotwords_config=hotwords_config_from_request
        )

        duration_view = round(time.perf_counter() - start_time_view, 2)

        if transcription_result.get("status") == "success":
            logger.info(
                f"speech_to_text_view: Transcription successful. RequestID: {request_id}, Duration: {duration_view}s"
            )
            # CORRECTED format_json_response call
            # The 'results' field expects a list of dictionaries.
            # For speech-to-text, we'll have one result item.
            response_results = [
                {
                    "original_name": audio_url,  # Or a more user-friendly name if available
                    "status": "success",
                    "message": "Speech to text conversion successful.",
                    "transcription": transcription_result.get("transcription"),
                    "raw_response_details": transcription_result.get(
                        "raw_response"
                    ),  # Optional, for debugging/client use
                }
            ]
            return format_json_response(
                results=response_results,
                merge_output=False,  # Not applicable to single audio transcription
                request_id=request_id,
                duration_seconds=duration_view,
            )
        else:
            error_msg = transcription_result.get(
                "message", "Unknown transcription error."
            )
            logger.error(
                f"speech_to_text_view: Transcription failed. RequestID: {request_id}, Error: {error_msg}, Duration: {duration_view}s"
            )
            # CORRECTED format_error_response call
            # Details from raw_response can be logged but not directly passed to format_error_response as a 'details' param.
            # We can include some part of it in the message if needed, or log it more extensively.
            if transcription_result.get("raw_response"):
                logger.error(
                    f"speech_to_text_view: Raw error response from DashScope: {transcription_result.get('raw_response')}. RequestID: {request_id}"
                )

            return format_error_response(
                message=f"Speech to text conversion failed: {error_msg}",
                merge_output=False,  # Not applicable
                http_status=500,  # Or a more specific error code if available
                request_id=request_id,
                # duration_seconds handled by format_error_response indirectly or needs explicit pass-through if desired
            )

    except json.JSONDecodeError:
        logger.error(
            f"speech_to_text_view: Invalid JSON in request body. RequestID: {request_id}",
            exc_info=True,
        )
        # CORRECTED format_error_response call
        return format_error_response(
            message="Invalid JSON format in request body.",
            merge_output=False,  # Not applicable
            http_status=400,
            request_id=request_id,
        )
    except Exception as e:
        logger.error(
            f"speech_to_text_view: An unexpected error occurred. RequestID: {request_id}. Error: {e}",
            exc_info=True,
        )
        # CORRECTED format_error_response call
        return format_error_response(
            message=f"An unexpected server error occurred: {str(e)}",
            merge_output=False,  # Not applicable
            http_status=500,
            request_id=request_id,
        )


# =====================================================================================
# NEW: Text to Speech (TTS) View
# =====================================================================================


def _extract_text_from_file(file_path):
    """Extracts text from txt, pdf, or docx file."""
    text = ""
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif file_extension == ".pdf":
        if pdfplumber is None:
            raise ImportError(
                "pdfplumber library is not available, cannot process PDF files for TTS."
            )
        with pdfplumber.open(file_path) as pdf:
            all_pages = [
                page.extract_text() for page in pdf.pages if page.extract_text()
            ]
            text = "\n".join(all_pages)
    elif file_extension == ".docx":
        doc = Document(file_path)
        all_paras = [para.text for para in doc.paragraphs]
        text = "\n".join(all_paras)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

    return text


def _chunk_text(text, max_length=1000):
    """Splits text into chunks, trying to preserve sentences."""
    if not text:
        return []

    chunks = []
    while len(text) > max_length:
        # Find a good split point (end of sentence, or last space)
        split_at = text.rfind("。", 0, max_length)
        if split_at == -1:
            split_at = text.rfind("！", 0, max_length)
        if split_at == -1:
            split_at = text.rfind("？", 0, max_length)
        if split_at == -1:
            split_at = text.rfind("\n", 0, max_length)
        if split_at == -1:
            split_at = text.rfind(" ", 0, max_length)
        if split_at == -1:  # If no good split point, just force split
            split_at = max_length - 1

        chunks.append(text[: split_at + 1].strip())
        text = text[split_at + 1 :].strip()

    if text:  # Add the last remaining part
        chunks.append(text)

    return chunks


@login_required
@require_POST
def text_to_speech_view(request):
    request_id = generate_request_id()
    logger.info(f"text_to_speech_view: Received request. RID: {request_id}")

    # --- Setup Directories and Paths ---
    username = request.user.username
    today_date_str = datetime.now().strftime("%Y%m%d")
    temp_files_to_clean = []

    try:
        user_upload_dir, user_converted_dir = ensure_user_directories(
            username, today_date_str
        )
        if not user_upload_dir:
            raise Exception("Failed to create user directories.")

        # --- Parse Input ---
        voice_model = request.POST.get("voice_model")
        text_input = request.POST.get("text_input")
        file_inputs = request.FILES.getlist("file_input")

        if not voice_model:
            return format_error_response(
                message="必须选择一个音色模型。",
                request_id=request_id,
                merge_output=False,
            )
        if not text_input and not file_inputs:
            return format_error_response(
                message="请输入文本或上传一个文件。",
                request_id=request_id,
                merge_output=False,
            )

        # --- Extract Text ---
        original_input_name = "文本输入"
        full_text = ""
        if file_inputs:
            logger.info(
                f"text_to_speech_view: Processing {len(file_inputs)} uploaded files. RID: {request_id}"
            )
            all_texts = []
            original_filenames = []
            for file_input in file_inputs:
                temp_file_path, _, _ = save_uploaded_file(
                    file_input, user_upload_dir, request_id
                )
                temp_files_to_clean.append(
                    temp_file_path
                )  # FIX: Append string, not dict
                original_filenames.append(file_input.name)
                try:
                    extracted_text = _extract_text_from_file(temp_file_path)
                    all_texts.append(extracted_text)
                    logger.info(
                        f"text_to_speech_view: Extracted {len(extracted_text)} chars from {file_input.name}. RID: {request_id}"
                    )
                except Exception as e_extract:
                    logger.error(
                        f"text_to_speech_view: Failed to extract text from {file_input.name}. RID: {request_id}, Error: {e_extract}",
                        exc_info=True,
                    )
                    return format_error_response(
                        message=f"从文件 '{file_input.name}' 提取文本失败: {e_extract}",
                        request_id=request_id,
                        merge_output=False,
                    )

            full_text = "\n".join(all_texts)
            original_input_name = ", ".join(original_filenames)
        else:
            full_text = text_input
            logger.info(
                f"text_to_speech_view: Processing text input. RID: {request_id}"
            )

        if not full_text.strip():
            return format_error_response(
                message="输入文本为空。", request_id=request_id, merge_output=False
            )

        # --- Process Text and Synthesize ---
        text_chunks = _chunk_text(full_text)
        logger.info(
            f"text_to_speech_view: Text split into {len(text_chunks)} chunks. RID: {request_id}"
        )
        audio_chunks_paths = []

        tts_script_path = os.path.join(
            settings.BASE_DIR, "converter", "text_to_voice.py"
        )
        if not os.path.exists(tts_script_path):
            return format_error_response(
                message="服务器配置错误：找不到语音合成脚本。",
                request_id=request_id,
                merge_output=False,
            )

        for i, chunk in enumerate(text_chunks):
            chunk_num = i + 1
            logger.info(
                f"text_to_speech_view: Processing chunk {chunk_num}/{len(text_chunks)}. RID: {request_id}"
            )

            chunk_input_txt_path = os.path.join(
                user_upload_dir, f"tts_chunk_{request_id}_{i}.txt"
            )
            chunk_output_wav_path = os.path.join(
                user_upload_dir, f"tts_chunk_audio_{request_id}_{i}.wav"
            )
            # FIX: Append strings, not dicts
            temp_files_to_clean.extend([chunk_input_txt_path, chunk_output_wav_path])

            with open(chunk_input_txt_path, "w", encoding="utf-8") as f:
                f.write(chunk)

            command = [
                sys.executable,
                tts_script_path,
                "--text_file_path",
                chunk_input_txt_path,
                "--voice",
                voice_model,
                "--output_file",
                chunk_output_wav_path,
            ]
            logger.info(
                f"text_to_speech_view: Executing TTS script for chunk {chunk_num}. Command: {' '.join(command)}. RID: {request_id}"
            )

            result = subprocess.run(
                command, capture_output=True, text=True, encoding="utf-8"
            )

            if result.returncode != 0:
                error_message = (
                    f"语音合成脚本执行失败 (片段 {chunk_num})。STDERR: {result.stderr}"
                )
                logger.error(
                    f"text_to_speech_view: TTS script failed. RID: {request_id}. {error_message}"
                )
                raise RuntimeError(error_message)

            if not os.path.exists(chunk_output_wav_path):
                error_message = "语音合成脚本未生成预期的音频文件。"
                logger.error(
                    f"text_to_speech_view: {error_message} Path: {chunk_output_wav_path}. RID: {request_id}"
                )
                raise RuntimeError(error_message)

            audio_chunks_paths.append(chunk_output_wav_path)

        # --- Concatenate Audio Chunks ---
        final_audio_path = ""
        if not audio_chunks_paths:
            raise RuntimeError("没有生成任何音频片段。")
        elif len(audio_chunks_paths) == 1:
            final_audio_path = audio_chunks_paths[0]
            logger.info(
                f"text_to_speech_view: Only one audio chunk created, no merge needed. RID: {request_id}"
            )
        else:
            logger.info(
                f"text_to_speech_view: Concatenating {len(audio_chunks_paths)} audio chunks. RID: {request_id}"
            )
            concat_list_path = os.path.join(
                user_upload_dir, f"tts_concat_list_{request_id}.txt"
            )
            temp_files_to_clean.append(concat_list_path)  # FIX: Append string, not dict

            with open(concat_list_path, "w", encoding="utf-8") as f:
                for chunk_path in audio_chunks_paths:
                    f.write(f"file '{chunk_path.replace(os.sep, '/')}'\n")

            concatenated_output_path = os.path.join(
                user_upload_dir, f"tts_merged_{request_id}.wav"
            )
            temp_files_to_clean.append(
                concatenated_output_path
            )  # FIX: Append string, not dict

            concat_command = [
                "ffmpeg",
                "-y",
                "-f",
                "concat",
                "-safe",
                "0",
                "-i",
                concat_list_path,
                "-c",
                "copy",
                concatenated_output_path,
            ]

            logger.info(
                f"text_to_speech_view: Executing ffmpeg concat command. RID: {request_id}"
            )
            result = subprocess.run(concat_command, capture_output=True, text=True)

            if result.returncode != 0:
                error_message = f"音频片段合并失败。FFMPEG STDERR: {result.stderr}"
                logger.error(
                    f"text_to_speech_view: FFMPEG failed. RID: {request_id}. {error_message}"
                )
                raise RuntimeError(error_message)

            final_audio_path = concatenated_output_path

        # --- Finalize and Prepare Response ---
        final_filename_display = f"tts_output_{request_id}.wav"
        final_destination_path = os.path.join(
            user_converted_dir, final_filename_display
        )

        logger.info(
            f"text_to_speech_view: Moving final audio from {final_audio_path} to {final_destination_path}. RID: {request_id}"
        )
        shutil.move(final_audio_path, final_destination_path)

        # FIX: Remove string path from cleanup list
        temp_files_to_clean = [p for p in temp_files_to_clean if p != final_audio_path]

        meta_file_path = f"{final_destination_path}.meta"
        with open(meta_file_path, "w", encoding="utf-8") as mf:
            mf.write(original_input_name)

        download_url = reverse(
            "converter:download_converted_file",
            args=[username, today_date_str, final_filename_display],
        )

        response_results = [
            {
                "original_name": original_input_name,
                "converted_name": final_filename_display,
                "download_url": download_url,
                "status": "success",
                "message": "文字转语音成功。",
            }
        ]

        logger.info(
            f"text_to_speech_view: Processing successful. Returning result. RID: {request_id}"
        )
        # FIX: Add missing merge_output argument
        return format_json_response(
            results=response_results, request_id=request_id, merge_output=False
        )

    except Exception as e:
        logger.error(
            f"text_to_speech_view: A critical error occurred. RID: {request_id}. Error: {e}",
            exc_info=True,
        )
        # FIX: Add missing merge_output argument
        return format_error_response(
            message=f"处理失败: {str(e)}", request_id=request_id, merge_output=False
        )
    finally:
        if temp_files_to_clean:
            logger.info(
                f"text_to_speech_view: Cleaning up {len(temp_files_to_clean)} temporary files. RID: {request_id}"
            )
            cleanup_temp_files(temp_files_to_clean, request_id)
