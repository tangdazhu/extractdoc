"""
DOCX exporter for OCR text extraction system.

This module handles exporting structured content to Microsoft Word DOCX format
with proper formatting, tables, and layout preservation.
"""

import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise ImportError(
        "python-docx library is required. Install with: pip install python-docx"
    )

from .base_exporter import BaseExporter, ExportError

logger = logging.getLogger(__name__)


class DocxExporter(BaseExporter):
    """DOCX document exporter with advanced formatting capabilities."""

    def __init__(self, output_path: str):
        """Initialize DOCX exporter."""
        super().__init__(output_path)
        self.document = None
        self._init_styles()

    def _init_styles(self):
        """Initialize document styles and formatting options."""
        self.styles = {
            "title": {"font_size": 16, "bold": True, "color": RGBColor(0, 0, 0)},
            "subtitle": {"font_size": 14, "bold": True, "color": RGBColor(64, 64, 64)},
            "heading1": {"font_size": 14, "bold": True, "color": RGBColor(0, 0, 0)},
            "heading2": {"font_size": 12, "bold": True, "color": RGBColor(32, 32, 32)},
            "normal": {"font_size": 11, "bold": False, "color": RGBColor(0, 0, 0)},
            "table_header": {
                "font_size": 10,
                "bold": True,
                "color": RGBColor(255, 255, 255),
                "background": RGBColor(68, 114, 196),
            },
            "table_cell": {"font_size": 10, "bold": False, "color": RGBColor(0, 0, 0)},
        }

    def get_export_format(self) -> str:
        """Get export format identifier."""
        return "docx"

    def export_document(
        self, document_structure: Dict[str, Any], metadata: Optional[Dict] = None
    ) -> bool:
        """
        Export complete document structure to DOCX.

        Args:
            document_structure: Structured document data
            metadata: Optional document metadata

        Returns:
            True if export successful
        """
        try:
            if not self.validate_output_path():
                return False

            # Prepare data for export
            prepared_data = self.prepare_export_data(document_structure)

            # Create new document
            self.document = Document()

            # Add document properties
            self._set_document_properties(prepared_data, metadata)

            # Add title if available
            title = prepared_data.get("title", "")
            if title:
                self._add_title(title)

            # Process sections
            sections = prepared_data.get("sections", [])
            for section in sections:
                self._add_section(section)

            # Save document
            self.document.save(str(self.output_path))
            self.logger.info(f"Document exported successfully to {self.output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error exporting document: {e}")
            raise ExportError(f"Failed to export DOCX document: {e}")

    def export_tables(
        self, tables: List[List[List[str]]], table_metadata: Optional[List[Dict]] = None
    ) -> bool:
        """
        Export tables to DOCX document.

        Args:
            tables: List of table structures
            table_metadata: Optional metadata for each table

        Returns:
            True if export successful
        """
        try:
            if not self.validate_output_path():
                return False

            # Create new document if not exists
            if not self.document:
                self.document = Document()

            # Add each table
            for i, table_data in enumerate(tables):
                if table_data:  # Only add non-empty tables
                    metadata = (
                        table_metadata[i]
                        if table_metadata and i < len(table_metadata)
                        else {}
                    )
                    self._add_table(table_data, metadata)

                    # Add spacing between tables
                    if i < len(tables) - 1:
                        self.document.add_paragraph()

            # Save document
            self.document.save(str(self.output_path))
            self.logger.info(f"Tables exported successfully to {self.output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error exporting tables: {e}")
            return False

    def export_text_content(self, text_contents: List[Dict]) -> bool:
        """
        Export text content to DOCX document.

        Args:
            text_contents: List of formatted text content

        Returns:
            True if export successful
        """
        try:
            if not self.validate_output_path():
                return False

            # Create new document if not exists
            if not self.document:
                self.document = Document()

            # Add each text content
            for content in text_contents:
                self._add_text_content(content)

            # Save document
            self.document.save(str(self.output_path))
            self.logger.info(
                f"Text content exported successfully to {self.output_path}"
            )
            return True

        except Exception as e:
            self.logger.error(f"Error exporting text content: {e}")
            return False

    def _set_document_properties(self, document_data: Dict, metadata: Optional[Dict]):
        """Set document properties and metadata."""
        core_properties = self.document.core_properties

        # Set basic properties
        core_properties.title = document_data.get("title", "OCR Extracted Document")
        core_properties.author = "OCR Text Extraction System"
        core_properties.subject = "Extracted text and tables from images"

        # Add metadata if available
        if metadata:
            if "created" in metadata:
                core_properties.created = metadata["created"]
            if "modified" in metadata:
                core_properties.modified = metadata["modified"]

    def _add_title(self, title: str):
        """Add document title with formatting."""
        title_paragraph = self.document.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply title formatting
        run = title_paragraph.runs[0]
        self._apply_text_formatting(run, self.styles["title"])

    def _add_section(self, section: Dict):
        """Add a document section with title and content."""
        title = section.get("title", "")
        level = section.get("level", 1)
        content = section.get("content", [])

        # Add section title
        if title:
            heading_level = min(level, 3)  # Limit to 3 levels
            section_heading = self.document.add_heading(title, level=heading_level)

            # Apply heading formatting
            if section_heading.runs:
                style_key = (
                    f"heading{heading_level}" if heading_level <= 2 else "normal"
                )
                self._apply_text_formatting(
                    section_heading.runs[0],
                    self.styles.get(style_key, self.styles["normal"]),
                )

        # Add section content
        for item in content:
            self._add_content_item(item)

    def _add_content_item(self, item: Dict):
        """Add a single content item to the document."""
        item_type = item.get("type", "text")
        content = item.get("content", "")

        if item_type == "text":
            self._add_text_content(item)
        elif item_type == "table":
            self._add_table(content, item.get("metadata", {}))
        elif item_type == "image":
            self._add_image_placeholder(content, item.get("metadata", {}))

    def _add_text_content(self, text_item: Dict):
        """Add formatted text content."""
        text = text_item.get("content", "")
        subtype = text_item.get("subtype", "paragraph")

        if not text.strip():
            return

        # Create paragraph
        paragraph = self.document.add_paragraph()

        # Determine formatting based on subtype
        if subtype == "title":
            paragraph.style = "Heading 1"
        elif subtype == "subtitle":
            paragraph.style = "Heading 2"
        elif subtype == "list_item":
            paragraph.style = "List Paragraph"
            # Add bullet point
            text = f"• {text}"
        else:
            paragraph.style = "Normal"

        # Add text with formatting
        run = paragraph.add_run(text)
        style_key = self._get_style_key_for_subtype(subtype)
        self._apply_text_formatting(
            run, self.styles.get(style_key, self.styles["normal"])
        )

    def _get_style_key_for_subtype(self, subtype: str) -> str:
        """Get style key for content subtype."""
        style_mapping = {
            "title": "title",
            "subtitle": "subtitle",
            "paragraph": "normal",
            "list_item": "normal",
            "header": "normal",
            "footer": "normal",
        }
        return style_mapping.get(subtype, "normal")

    def _add_table(self, table_data: List[List[str]], metadata: Optional[Dict] = None):
        """Add a formatted table to the document."""
        if not table_data:
            return

        # Determine table dimensions
        max_cols = max(len(row) for row in table_data) if table_data else 0
        if max_cols == 0:
            return

        # Create table
        table = self.document.add_table(rows=len(table_data), cols=max_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        self._set_table_column_widths(table, max_cols)

        # Fill table data
        for row_idx, row_data in enumerate(table_data):
            table_row = table.rows[row_idx]

            for col_idx, cell_data in enumerate(row_data):
                if col_idx < max_cols:
                    cell = table_row.cells[col_idx]
                    cell.text = str(cell_data).strip()

                    # Apply cell formatting
                    self._format_table_cell(cell, is_header=(row_idx == 0))

        # Add spacing after table
        self.document.add_paragraph()

    def _set_table_column_widths(self, table, num_cols: int):
        """Set optimal column widths for the table."""
        # Calculate column width based on page width
        page_width = Inches(8.5) - Inches(2)  # Letter size minus margins
        col_width = page_width / num_cols

        for col in table.columns:
            col.width = col_width

    def _format_table_cell(self, cell, is_header: bool = False):
        """Format a single table cell."""
        # Set cell text formatting
        if cell.paragraphs:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            )

            if paragraph.runs:
                run = paragraph.runs[0]
                style_key = "table_header" if is_header else "table_cell"
                self._apply_text_formatting(run, self.styles[style_key])

        # Set cell background for headers
        if is_header:
            self._set_cell_background(cell, self.styles["table_header"]["background"])

    def _set_cell_background(self, cell, color: RGBColor):
        """Set cell background color."""
        try:
            # Access the cell's XML element
            cell_xml = cell._tc
            cell_properties = cell_xml.get_or_add_tcPr()

            # Create shading element
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), f"{color.r:02x}{color.g:02x}{color.b:02x}")
            cell_properties.append(shading)
        except Exception as e:
            self.logger.warning(f"Could not set cell background: {e}")

    def _apply_text_formatting(self, run, style: Dict):
        """Apply text formatting to a run."""
        if "font_size" in style:
            run.font.size = Pt(style["font_size"])

        if "bold" in style:
            run.font.bold = style["bold"]

        if "color" in style:
            run.font.color.rgb = style["color"]

        if "italic" in style:
            run.font.italic = style["italic"]

    def _add_image_placeholder(self, image_data: Dict, metadata: Optional[Dict] = None):
        """Add image placeholder text."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run("[Image placeholder]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Add image metadata if available
        if metadata:
            image_info = f" - {metadata.get('filename', 'Unknown file')}"
            paragraph.add_run(image_info)

    def add_page_break(self):
        """Add a page break to the document."""
        if self.document:
            self.document.add_page_break()

    def set_page_margins(
        self,
        top: float = 1.0,
        bottom: float = 1.0,
        left: float = 1.0,
        right: float = 1.0,
    ):
        """
        Set page margins in inches.

        Args:
            top: Top margin in inches
            bottom: Bottom margin in inches
            left: Left margin in inches
            right: Right margin in inches
        """
        if self.document:
            sections = self.document.sections
            for section in sections:
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)

    def add_header_footer(self, header_text: str = "", footer_text: str = ""):
        """Add header and footer to the document."""
        if not self.document:
            return

        section = self.document.sections[0]

        # Add header
        if header_text:
            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = header_text
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add footer
        if footer_text:
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = footer_text
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
