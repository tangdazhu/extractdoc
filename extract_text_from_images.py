#!/usr/bin/env python3
"""
OCR文本提取系统 - 极简还原版
严格对齐 extract_text_from_images_original.py，只保留特殊表格和普通文本两种处理。
"""
import os
import sys
import argparse
import logging
from pathlib import Path
import traceback

if sys.platform.startswith("win"):
    import io

    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")
    try:
        os.system("chcp 65001 > nul")
    except:
        pass
current_dir = Path(__file__).parent
sys.path.insert(0, str(current_dir))


def safe_print(text, fallback_text=None):
    try:
        print(text)
    except UnicodeEncodeError:
        if fallback_text:
            print(fallback_text)
        else:
            print(text.encode("ascii", "replace").decode("ascii"))


try:
    from core.ocr_engine import OCREngine
    from handlers.special_tables import SpecialTableHandler
    from utils.validation import ValidationUtils
    from config.settings import Settings
    from core.table_detector import (
        reconstruct_table_from_coordinates,
        add_reconstructed_table_to_docx,
        process_mixed_table_text_content,
    )
    from extract_text_from_images_original import (
        add_formatted_content_to_pptx,
        segment_text,
    )

    MODULES_AVAILABLE = True
except ImportError as e:
    safe_print(f"✗ 模块化组件导入失败: {e}", f"✗ Module import failed: {e}")
    traceback.print_exc()
    safe_print(
        "如果您需要使用原始版本，请使用: extract_text_from_images_original.py",
        "If you need the original version, use: extract_text_from_images_original.py",
    )
    MODULES_AVAILABLE = False


class ModularOCRSystem:
    def __init__(self, config_path=None):
        self.settings = Settings(config_path)
        self.logger = self._setup_logging()
        self.ocr_engine = None
        self.special_handler = SpecialTableHandler()
        self.logger.info("模块化OCR系统初始化完成")

    def _setup_logging(self):
        logger = logging.getLogger("ocr_system")
        logger.setLevel(logging.DEBUG)
        # 彻底移除所有旧的 handler
        for handler in logger.handlers[:]:
            logger.removeHandler(handler)
            handler.close()
        log_file = self.settings.get("log_file", "app.log")
        if not log_file:
            log_file = "app.log"
        file_handler = logging.FileHandler(log_file, encoding="utf-8")
        file_formatter = logging.Formatter(
            "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
        )
        file_handler.setFormatter(file_formatter)
        file_handler.setLevel(logging.DEBUG)
        logger.addHandler(file_handler)
        console_handler = logging.StreamHandler()
        console_formatter = logging.Formatter("%(levelname)s: %(message)s")
        console_handler.setFormatter(console_formatter)
        console_handler.setLevel(logging.INFO)
        logger.addHandler(console_handler)
        logger.propagate = False  # 禁止向上冒泡，防止被 root logger 过滤
        logger.info(f"Logger file path: {log_file}")
        logger.debug("[DEBUG-INIT] Only my handlers should exist now.")
        return logger

    def _initialize_ocr(self):
        if self.ocr_engine is None:
            self.ocr_engine = OCREngine()
            self.logger.info("OCR引擎初始化完成")

    def process_image(
        self, image_path, output_path=None, output_format="docx", content_format="auto"
    ):
        self.logger.debug(
            "!!! DEBUG TEST: This is a debug log, should appear in app.log !!!"
        )
        if output_path is None:
            base_name = os.path.splitext(os.path.basename(image_path))[0]
            output_path = f"{base_name}_extracted.{output_format}"
        from docx import Document

        doc = Document()
        try:
            self.logger.info(f"开始处理图片: {image_path}")
            self.logger.debug(
                f"process_image: input_path={image_path}, output_path={output_path}, output_format={output_format}, content_format={content_format}"
            )
            if not ValidationUtils.validate_image_file(image_path):
                raise ValueError(f"无效的图片路径: {image_path}")
            self._initialize_ocr()
            layout_elements = self.ocr_engine.extract_layout_elements(image_path)
            self.logger.debug(
                f"process_image: layout_elements type={type(layout_elements)}, count={len(layout_elements) if layout_elements else 0}"
            )
            filename = os.path.basename(image_path)
            mixed_content_result = None
            if layout_elements:
                mixed_content_result = process_mixed_table_text_content(
                    layout_elements, self.logger
                )
            title_texts = []
            if mixed_content_result and len(mixed_content_result) == 3:
                _, _, title_texts = mixed_content_result
            if title_texts:
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                for title in title_texts:
                    if title.strip():
                        para = doc.add_paragraph(title.strip())
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(f"OCR提取结果 - {filename}", level=1)
            if self.special_handler.is_special_table(filename):
                self.logger.info(f"检测到特殊表格，使用专用处理器: {filename}")
                self.logger.debug(
                    f"process_image: special_handler branch, filename={filename}, layout_elements={layout_elements}"
                )
                self.special_handler.process_special_table(
                    filename, layout_elements, doc
                )
            else:
                if not layout_elements:
                    self.logger.warning(f"布局分析失败，尝试简单文本提取: {image_path}")
                    simple_text = self.ocr_engine.extract_text_simple(image_path)
                    self.logger.debug(f"process_image: simple_text={simple_text}")
                    if simple_text and simple_text != "No text detected in this image.":
                        for line in simple_text.split("\n"):
                            if line.strip():
                                doc.add_paragraph(line)
                    else:
                        self.logger.warning(f"未能从图片中提取到任何内容: {image_path}")
                        doc.add_paragraph(
                            f"[No content could be extracted from {filename}]"
                        )
                else:
                    if mixed_content_result:
                        table_rows, remaining_elements, _ = mixed_content_result
                        add_reconstructed_table_to_docx(doc, table_rows)
                        if "ppt" in filename.lower() or filename.lower().endswith(
                            (".ppt", ".pptx")
                        ):
                            self.logger.info(
                                f"已识别为PPT表格，跳过文本内容: {filename}"
                            )
                            doc.save(output_path)
                            self.logger.info(f"处理完成，文档已保存: {output_path}")
                            return output_path
                    else:
                        is_table_detected, table_rows = (
                            reconstruct_table_from_coordinates(
                                layout_elements, self.logger
                            )
                        )
                        if is_table_detected:
                            self.logger.info(f"通过坐标重建表格结构: {filename}")
                            add_reconstructed_table_to_docx(doc, table_rows)
                            if "ppt" in filename.lower() or filename.lower().endswith(
                                (".ppt", ".pptx")
                            ):
                                self.logger.info(
                                    f"已识别为PPT表格，跳过文本内容: {filename}"
                                )
                                doc.save(output_path)
                                self.logger.info(f"处理完成，文档已保存: {output_path}")
                                return output_path
                        else:
                            self.logger.info(
                                f"未检测到表格结构，按段落处理: {filename}"
                            )
                            for e in layout_elements:
                                if (
                                    isinstance(e, (list, tuple))
                                    and len(e) > 1
                                    and isinstance(e[1], (list, tuple))
                                    and len(e[1]) > 0
                                ):
                                    text = str(e[1][0])
                                elif isinstance(e, dict) and "text" in e:
                                    text = str(e["text"])
                                else:
                                    text = None
                                if text and text.strip():
                                    doc.add_paragraph(text.strip())
            doc.save(output_path)
            self.logger.info(f"处理完成，文档已保存: {output_path}")
            return output_path
        except Exception as e:
            self.logger.error(f"处理图片时发生错误: {e}", exc_info=True)
            from docx import Document

            error_doc = Document()
            error_doc.add_paragraph(
                f"处理图片 '{os.path.basename(image_path)}' 时发生错误。"
            )
            error_doc.add_paragraph(str(e))
            error_doc.save(output_path)
            return output_path

    def _get_image_files(self, directory):
        image_extensions = [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]
        image_files = []
        for ext in image_extensions:
            pattern = os.path.join(directory, f"*{ext}")
            image_files.extend(glob.glob(pattern))
        return sorted(image_files)

    def process_directory(
        self, input_dir, output_path=None, output_format="docx", content_format="auto"
    ):
        try:
            self.logger.info(f"开始批量处理目录: {input_dir}")
            self.logger.debug(
                f"process_directory: input_dir={input_dir}, output_path={output_path}, output_format={output_format}, content_format={content_format}"
            )
            image_files = self._get_image_files(input_dir)
            self.logger.debug(
                f"process_directory: found {len(image_files)} image files: {image_files}"
            )
            if not image_files:
                self.logger.warning(f"目录中未找到图片文件: {input_dir}")
                return None
            self.logger.info(f"找到 {len(image_files)} 个图片文件")
            from docx import Document

            merged_doc = Document()
            for i, image_path in enumerate(image_files):
                filename = os.path.basename(image_path)
                self.logger.info(f"处理进度: {i+1}/{len(image_files)} - {filename}")
                self.logger.debug(f"process_directory: processing {filename}")
                layout_elements = self.ocr_engine.extract_layout_elements(image_path)
                mixed_content_result = None
                if layout_elements:
                    mixed_content_result = process_mixed_table_text_content(
                        layout_elements, self.logger
                    )
                title_texts = []
                if mixed_content_result and len(mixed_content_result) == 3:
                    _, _, title_texts = mixed_content_result
                if title_texts:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    for title in title_texts:
                        if title.strip():
                            para = merged_doc.add_paragraph(title.strip())
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    merged_doc.add_heading(f"文件: {filename}", level=1)
                try:
                    self._initialize_ocr()
                    is_ppt_style = False
                    all_text_lines = []
                    if self.special_handler.is_special_table(filename):
                        self.logger.info(f"检测到特殊表格，使用专用处理器: {filename}")
                        self.logger.debug(
                            f"process_directory: special_handler branch, filename={filename}, layout_elements={layout_elements}"
                        )
                        self.special_handler.process_special_table(
                            filename, layout_elements, merged_doc
                        )
                    else:
                        if not layout_elements:
                            self.logger.warning(
                                f"布局分析失败，尝试简单文本提取: {filename}"
                            )
                            simple_text = self.ocr_engine.extract_text_simple(
                                image_path
                            )
                            self.logger.debug(
                                f"process_directory: simple_text={simple_text}"
                            )
                            if (
                                simple_text
                                and simple_text != "No text detected in this image."
                            ):
                                for line in simple_text.split("\n"):
                                    if line.strip():
                                        merged_doc.add_paragraph(line)
                        else:
                            if mixed_content_result:
                                table_rows, remaining_elements, _ = mixed_content_result
                                # 检查是否为PPT风格
                                formatted_content = segment_text(
                                    "\n".join(
                                        [
                                            item
                                            for row in table_rows
                                            for item in row
                                            if isinstance(item, str)
                                        ]
                                    )
                                )
                                has_title = any(
                                    item.get("type") == "title"
                                    for item in formatted_content
                                )
                                has_main_sections = any(
                                    item.get("type") == "numbered_main"
                                    for item in formatted_content
                                )
                                has_bullets = any(
                                    item.get("type") in ["bullet_sub", "bullet"]
                                    for item in formatted_content
                                )
                                if content_format == "ppt":
                                    is_ppt_style = True
                                elif content_format == "auto":
                                    if has_title and has_main_sections and has_bullets:
                                        is_ppt_style = True
                                if is_ppt_style:
                                    # 只处理当前图片内容，每次只插入一页slide
                                    def single_slide_content(formatted_content):
                                        from docx.shared import Pt
                                        from docx.enum.text import WD_ALIGN_PARAGRAPH

                                        if formatted_content:
                                            title_item = next(
                                                (
                                                    item
                                                    for item in formatted_content
                                                    if item.get("type") == "title"
                                                ),
                                                None,
                                            )
                                            if title_item:
                                                para = merged_doc.add_heading(
                                                    title_item["text"], level=1
                                                )
                                                para.alignment = (
                                                    WD_ALIGN_PARAGRAPH.CENTER
                                                )
                                                run = (
                                                    para.runs[0]
                                                    if para.runs
                                                    else para.add_run(
                                                        title_item["text"]
                                                    )
                                                )
                                                run.font.size = Pt(18)
                                                run.font.bold = True
                                            for item in formatted_content:
                                                if item.get("type") == "title":
                                                    continue
                                                text = item.get("text", "")
                                                if item.get("type") == "subtitle":
                                                    para = merged_doc.add_paragraph(
                                                        text
                                                    )
                                                    para.alignment = (
                                                        WD_ALIGN_PARAGRAPH.CENTER
                                                    )
                                                    run = (
                                                        para.runs[0]
                                                        if para.runs
                                                        else para.add_run(text)
                                                    )
                                                    run.font.size = Pt(14)
                                                    run.font.italic = True
                                                elif (
                                                    item.get("type") == "section_header"
                                                ):
                                                    para = merged_doc.add_paragraph()
                                                    run = para.add_run(text)
                                                    run.font.size = Pt(14)
                                                    run.font.bold = True
                                                elif item.get("type") in [
                                                    "bullet_sub",
                                                    "bullet",
                                                    "numbered_paren",
                                                    "section_sub",
                                                    "text_sub",
                                                ]:
                                                    para = merged_doc.add_paragraph()
                                                    para.paragraph_format.left_indent = Pt(
                                                        36
                                                    )
                                                    bullet_run = para.add_run("• ")
                                                    bullet_run.font.size = Pt(12)
                                                    text_run = para.add_run(text)
                                                    text_run.font.size = Pt(12)
                                                else:
                                                    para = merged_doc.add_paragraph(
                                                        text
                                                    )
                                                    run = (
                                                        para.runs[0]
                                                        if para.runs
                                                        else para.add_run(text)
                                                    )
                                                    run.font.size = Pt(12)

                                    single_slide_content(formatted_content)
                                else:
                                    add_reconstructed_table_to_docx(
                                        merged_doc, table_rows
                                    )
                            else:
                                is_table_detected, table_rows = (
                                    reconstruct_table_from_coordinates(
                                        layout_elements, self.logger
                                    )
                                )
                                if is_table_detected:
                                    add_reconstructed_table_to_docx(
                                        merged_doc, table_rows
                                    )
                                else:
                                    for e in layout_elements:
                                        if (
                                            isinstance(e, (list, tuple))
                                            and len(e) > 1
                                            and isinstance(e[1], (list, tuple))
                                            and len(e[1]) > 0
                                        ):
                                            text = str(e[1][0])
                                        elif isinstance(e, dict) and "text" in e:
                                            text = str(e["text"])
                                        else:
                                            text = None
                                        if text and text.strip():
                                            merged_doc.add_paragraph(text.strip())
                except Exception as e:
                    self.logger.error(f"处理文件 {image_path} 时出错: {e}")
                    merged_doc.add_paragraph(
                        f"--- ERROR PROCESSING {filename}: {e} ---"
                    )
                # 只要不是最后一张图片，都插入分页符
                if i < len(image_files) - 1:
                    merged_doc.add_page_break()
            if output_path is None:
                output_path = os.path.join(input_dir, f"extracted_text.{output_format}")
            merged_doc.save(output_path)
            self.logger.info(f"批量处理完成: {output_path}")
            return output_path
        except Exception as e:
            self.logger.error(f"批量处理时发生错误: {e}", exc_info=True)
            raise


def main(
    input_path_arg=None,
    output_path_arg=None,
    output_format_arg="docx",
    content_format="auto",
):
    if not MODULES_AVAILABLE:
        print("\n" + "=" * 60)
        safe_print("错误: 模块化组件不可用", "Error: Modular components not available")
        print("=" * 60)
        safe_print("可能的解决方案:", "Possible solutions:")
        safe_print(
            "1. 确保所有模块文件已正确创建",
            "1. Ensure all module files are correctly created",
        )
        safe_print(
            "2. 检查Python环境和依赖", "2. Check Python environment and dependencies"
        )
        safe_print(
            "3. 使用原始版本: extract_text_from_images_original.py",
            "3. Use original version: extract_text_from_images_original.py",
        )
        print("=" * 60)
        return
    print("\n" + "=" * 60)
    safe_print(
        "📁 OCR文本提取系统 - 极简还原版",
        "📁 OCR Text Extraction System - Minimal Restored Version",
    )
    print("=" * 60)
    try:
        ocr_system = ModularOCRSystem()
        if input_path_arg:
            if os.path.isfile(input_path_arg):
                safe_print(
                    f"🖼️  处理单个图片: {input_path_arg}",
                    f"🖼️  Processing single image: {input_path_arg}",
                )
                result = ocr_system.process_image(
                    input_path_arg, output_path_arg, output_format_arg, content_format
                )
            elif os.path.isdir(input_path_arg):
                safe_print(
                    f"📁 批量处理目录: {input_path_arg}",
                    f"📁 Batch processing directory: {input_path_arg}",
                )
                result = ocr_system.process_directory(
                    input_path_arg, output_path_arg, output_format_arg, content_format
                )
            else:
                safe_print(
                    f"❌ 错误: 路径不存在 - {input_path_arg}",
                    f"❌ Error: Path does not exist - {input_path_arg}",
                )
                return
        else:
            default_dir = ocr_system.settings.get("input_dir", ".")
            safe_print(
                f"📁 使用默认目录: {default_dir}",
                f"📁 Using default directory: {default_dir}",
            )
            result = ocr_system.process_directory(
                default_dir, output_path_arg, output_format_arg, content_format
            )
        if result:
            safe_print(
                f"\n🎉 处理完成! 输出文件: {result}",
                f"\n🎉 Processing complete! Output file: {result}",
            )
            safe_print(
                "✨ 极简还原成功！系统运行正常。",
                "✨ Minimal restoration successful! System running normally.",
            )
        else:
            safe_print(
                "\n⚠️  处理完成，但未生成输出文件",
                "\n⚠️  Processing complete, but no output file generated",
            )
    except Exception as e:
        safe_print(
            f"\n❌ 处理过程中发生错误: {e}",
            f"\n❌ Error occurred during processing: {e}",
        )
        logging.error(f"主函数错误: {e}", exc_info=True)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Extract text and tables from images to DOCX or PDF. (Minimal restored version)"
    )
    parser.add_argument(
        "input_path",
        nargs="?",
        default=None,
        help="Path to a single input image file or directory.",
    )
    parser.add_argument(
        "output_path",
        nargs="?",
        default=None,
        help="Path for the output file (e.g., document.docx or document.pdf).",
    )
    parser.add_argument(
        "--format",
        choices=["docx", "pdf"],
        default="docx",
        help="Output format (docx or pdf). Default is docx.",
    )
    parser.add_argument(
        "--content-format",
        choices=["auto", "docx", "ppt"],
        default="auto",
        help="Content formatting style: auto (detect), docx (document style), ppt (slide style). Default is auto.",
    )
    args = parser.parse_args()
    main(
        input_path_arg=args.input_path,
        output_path_arg=args.output_path,
        output_format_arg=args.format,
        content_format=args.content_format,
    )
