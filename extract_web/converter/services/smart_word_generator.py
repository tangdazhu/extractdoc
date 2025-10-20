# -*- coding: utf-8 -*-
"""
智能Word文档生成器

基于AI分析的PDF结构，生成格式化的Word文档，包含：
- 标题页
- 结构化内容
- 表格嵌入
- 图片嵌入
- 智能分页
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# 导入配置管理器
import sys
sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))
from utils.config_manager import config

logger = logging.getLogger("converter")


class SmartWordGenerator:
    """智能Word文档生成器"""
    
    def __init__(self, style_config: Optional[Dict] = None):
        """
        初始化生成器
        
        Args:
            style_config: 样式配置，如果为None则从配置文件加载
        """
        self.doc = None
        # 从配置文件加载，允许运行时覆盖
        if style_config is None:
            self.config = config.get_section("word_generation")
        else:
            self.config = style_config
        
    def generate_word(
        self,
        document_structure: Dict,
        page_analyses: List[Dict],
        multimodal_data: Dict,
        request_id: str
    ) -> Document:
        """
        生成智能Word文档
        
        Args:
            document_structure: 文档整体结构分析
            page_analyses: 每页的详细分析
            multimodal_data: 多模态数据(文本、表格、图片)
            request_id: 请求ID
            
        Returns:
            生成的Document对象
        """
        logger.info("开始智能Word生成(新版),RequestID=%s", request_id)
        
        # 创建新文档
        self.doc = Document()
        
        # 设置页边距
        self._set_page_margins()
        
        # 1. 创建标题页
        self._create_title_page(document_structure, multimodal_data)
        
        # 2. 创建内容页
        for page_analysis in page_analyses:
            self._create_content_page(page_analysis, multimodal_data, request_id)
        
        logger.info("智能Word生成完成,RequestID=%s", request_id)
        return self.doc
    
    def _set_page_margins(self):
        """设置页边距"""
        section = self.doc.sections[0]
        margins = self.config["page_margins"]
        section.top_margin = Inches(margins["top"])
        section.bottom_margin = Inches(margins["bottom"])
        section.left_margin = Inches(margins["left"])
        section.right_margin = Inches(margins["right"])
    
    def _create_title_page(self, document_structure: Dict, multimodal_data: Dict):
        """
        创建标题页
        
        Args:
            document_structure: 文档结构
            multimodal_data: 多模态数据
        """
        # 从document_structure中提取标题和副标题
        title_page = document_structure.get("title_page", {})
        elements = title_page.get("elements", {})
        title = elements.get("title", "")
        if not title:
            raise ValueError("未能从文档结构中提取到标题")
        subtitle = elements.get("subtitle", "")
        
        # 添加标题(居中、加粗、大字号)
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(self.config["font_sizes"]["title"])
        title_run.font.bold = True
        title_color = self.config["colors"]["title"]
        title_run.font.color.rgb = RGBColor(*title_color)
        
        # 添加副标题(如果有)
        if subtitle:
            subtitle_para = self.doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(self.config["font_sizes"]["subtitle"])
            subtitle_color = self.config["colors"]["subtitle"]
            subtitle_run.font.color.rgb = RGBColor(*subtitle_color)
        
        # 添加空行
        self.doc.add_paragraph()
        
        # 查找第一页的元数据表格(如果有)
        # 使用AI分析结果中的元数据信息
        metadata = elements.get("metadata", {})
        if metadata:
            # 将元数据转换为表格格式
            metadata_table = [[k, str(v)] for k, v in metadata.items() if v]
            if metadata_table:
                self._add_table(metadata_table, is_metadata=True)
        else:
            # 降级方案：查找第一页的小表格
            first_page_tables = [t for t in multimodal_data.get("tables", []) if t["page"] == 1]
            if first_page_tables:
                for table_info in first_page_tables:
                    table_data = table_info["data"]
                    # 小表格可能是元数据（行数少且列数少）
                    if len(table_data) <= 5 and len(table_data[0]) <= 3:
                        self._add_table(table_data, is_metadata=True)
                        break
        
        # 添加分页符
        self.doc.add_page_break()
        
        logger.info("已创建标题页: %s", title)
    
    def _create_content_page(
        self,
        page_analysis: Dict,
        multimodal_data: Dict,
        request_id: str
    ):
        """
        创建内容页
        
        Args:
            page_analysis: 页面分析结果
            multimodal_data: 多模态数据
            request_id: 请求ID
        """
        page_num = page_analysis.get("page_number", 0)
        page_title = page_analysis.get("title", "")
        layout_type = page_analysis.get("suggested_layout", "text_only")
        
        logger.debug("创建第%d页: %s, 布局=%s", page_num, page_title, layout_type)
        
        # 添加页面标题
        if page_title and page_num > 1:  # 跳过标题页
            heading = self.doc.add_heading(page_title, level=1)
            heading.runs[0].font.size = Pt(self.config["font_sizes"]["heading"])
            title_color = self.config["colors"]["title"]
            heading.runs[0].font.color.rgb = RGBColor(*title_color)
        
        # 第1页是标题页，已在_create_title_page中处理，这里跳过
        if page_num == 1:
            logger.debug("第1页是标题页，跳过内容渲染")
            return
        
        # 根据布局类型渲染内容
        if layout_type == "title_and_table":
            self._render_table_layout(page_num, page_analysis, multimodal_data)
        elif layout_type == "title_and_image":
            self._render_image_layout(page_num, page_analysis, multimodal_data)
        elif layout_type == "title_and_text":
            self._render_text_layout(page_num, page_analysis, multimodal_data)
        elif layout_type == "mixed":
            self._render_mixed_layout(page_num, page_analysis, multimodal_data)
        else:
            # 默认文本布局
            self._render_text_layout(page_num, page_analysis, multimodal_data)
        
        # 如果不是最后一页,添加分页符
        # (这里简化处理,实际可以根据内容长度智能决定是否分页)
        if page_num < len(multimodal_data.get("pages", [])):
            self.doc.add_page_break()
        
        logger.debug("已创建内容页: 第%d页 - %s", page_num, page_title)
    
    def _render_table_layout(
        self,
        page_num: int,
        page_analysis: Dict,
        multimodal_data: Dict
    ):
        """渲染表格布局"""
        # 获取该页的表格
        page_tables = [t for t in multimodal_data.get("tables", []) if t["page"] == page_num]
        
        if page_tables:
            for table_info in page_tables:
                table_data = table_info["data"]
                # 跳过元数据表格(已在标题页显示)
                if page_num == 1 and len(table_data) <= 5:
                    continue
                self._add_table(table_data)
                self.doc.add_paragraph()  # 表格后添加空行
    
    def _render_image_layout(
        self,
        page_num: int,
        page_analysis: Dict,
        multimodal_data: Dict
    ):
        """渲染图片布局"""
        # 添加AI重新组织的文本(如果有)
        formatted_content = page_analysis.get("formatted_content", "")
        if formatted_content:
            self._add_formatted_text(formatted_content)
            self.doc.add_paragraph()
        
        # 获取该页的图片
        page_images = [img for img in multimodal_data.get("images", []) if img["page"] == page_num]
        
        if page_images:
            for img_info in page_images:
                img_path = img_info["path"]
                self._add_image(img_path)
                self.doc.add_paragraph()  # 图片后添加空行
    
    def _render_text_layout(
        self,
        page_num: int,
        page_analysis: Dict,
        multimodal_data: Dict
    ):
        """渲染纯文本布局"""
        # 使用AI重新组织的文本
        formatted_content = page_analysis.get("formatted_content", "")
        if formatted_content:
            self._add_formatted_text(formatted_content)
        else:
            # 降级: 使用原始文本
            page_data = next((p for p in multimodal_data.get("pages", []) if p["page"] == page_num), None)
            if page_data:
                raw_text = page_data.get("text", "")
                if raw_text:
                    self._add_formatted_text(raw_text)
    
    def _render_mixed_layout(
        self,
        page_num: int,
        page_analysis: Dict,
        multimodal_data: Dict
    ):
        """渲染混合布局(文本+表格+图片)"""
        # 1. 添加文本
        formatted_content = page_analysis.get("formatted_content", "")
        if formatted_content:
            self._add_formatted_text(formatted_content)
            self.doc.add_paragraph()
        
        # 2. 添加表格
        page_tables = [t for t in multimodal_data.get("tables", []) if t["page"] == page_num]
        if page_tables:
            for table_info in page_tables:
                self._add_table(table_info["data"])
                self.doc.add_paragraph()
        
        # 3. 添加图片
        page_images = [img for img in multimodal_data.get("images", []) if img["page"] == page_num]
        if page_images:
            for img_info in page_images:
                self._add_image(img_info["path"])
                self.doc.add_paragraph()
    
    def _add_formatted_text(self, text: str):
        """
        添加格式化文本
        
        支持:
        - 列表项(以数字或符号开头)
        - 段落
        """
        if not text or not text.strip():
            return
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测列表项
            body_size = self.config["font_sizes"]["body"]
            if self._is_list_item(line):
                para = self.doc.add_paragraph(line, style='List Bullet')
                if para.runs:
                    para.runs[0].font.size = Pt(body_size)
            else:
                # 普通段落
                para = self.doc.add_paragraph(line)
                if para.runs:
                    para.runs[0].font.size = Pt(body_size)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _is_list_item(self, line: str) -> bool:
        """判断是否为列表项"""
        import re
        # 匹配: "1. ", "• ", "- ", "* " 等
        pattern = r'^(\d+[\.\)、]|[•\-\*])\s+'
        return bool(re.match(pattern, line))
    
    def _add_table(self, table_data: List[List[str]], is_metadata: bool = False):
        """
        添加表格
        
        Args:
            table_data: 表格数据
            is_metadata: 是否为元数据表格(样式不同)
        """
        if not table_data or len(table_data) == 0:
            return
        
        # 规范化表格(确保所有行列数一致)
        max_cols = max(len(row) for row in table_data)
        normalized_table = []
        for row in table_data:
            if len(row) < max_cols:
                normalized_row = row + [''] * (max_cols - len(row))
            else:
                normalized_row = row
            normalized_table.append(normalized_row)
        
        # 创建表格
        table = self.doc.add_table(rows=len(normalized_table), cols=max_cols)
        table.style = self.config["table"]["style"]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 填充数据
        for row_idx, row_data in enumerate(normalized_table):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_data).strip()
                
                # 格式化单元格
                table_font_size = self.config["font_sizes"]["table"]
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    if row_idx == 0:
                        # 表头: 加粗、居中
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if para.runs:
                            para.runs[0].font.bold = True
                            para.runs[0].font.size = Pt(table_font_size)
                    else:
                        # 数据行
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if para.runs:
                            para.runs[0].font.size = Pt(table_font_size)
        
        logger.debug("已添加表格: %d行x%d列", len(normalized_table), max_cols)
    
    def _add_image(self, img_path: Path):
        """
        添加图片
        
        Args:
            img_path: 图片路径
        """
        if not img_path or not img_path.exists():
            logger.warning("图片不存在: %s", img_path)
            return
        
        try:
            # 获取图片尺寸
            with Image.open(img_path) as img:
                width, height = img.size
            
            # 计算合适的显示尺寸(保持宽高比)
            img_config = self.config["image"]
            max_width_inches = img_config["max_width_inches"]
            max_height_inches = img_config["max_height_inches"]
            default_dpi = img_config["default_dpi"]
            
            # 将像素转换为英寸（使用配置的DPI）
            img_width_inches = width / float(default_dpi)
            img_height_inches = height / float(default_dpi)
            
            # 计算缩放比例
            width_ratio = max_width_inches / img_width_inches
            height_ratio = max_height_inches / img_height_inches
            scale_ratio = min(width_ratio, height_ratio, 1.0)  # 不放大
            
            # 计算最终尺寸
            final_width_inches = img_width_inches * scale_ratio
            final_height_inches = img_height_inches * scale_ratio
            
            # 转换为EMU单位
            final_width = Inches(final_width_inches)
            final_height = Inches(final_height_inches)
            
            # 添加图片
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(img_path), width=final_width, height=final_height)
            
            logger.debug("已添加图片: %s (%.1f x %.1f 英寸)", img_path.name, final_width_inches, final_height_inches)
            
        except Exception as e:
            logger.error("添加图片失败: %s, 错误: %s", img_path, e)
            # 添加占位符
            para = self.doc.add_paragraph()
            run = para.add_run(f"[图片: {img_path.name}]")
            run.font.italic = True
            placeholder_color = self.config["colors"]["placeholder"]
            run.font.color.rgb = RGBColor(*placeholder_color)
